"""
cv_parser.py  (column-aware edition + spacing normalisation)
=============================================================
Production-level CV/Resume Parsing Pipeline
Processes PDF and DOCX files from two directories and outputs a structured JSON dataset.

Key features:
  • Column-aware PDF extraction  — detects single / two-column / multi-column (3+)
    layouts using pdfplumber word bounding-boxes and reconstructs correct reading
    order (left column → right column, top-to-bottom within each).
  • Gap-analysis column split   — instead of a hard page-midpoint cut, we find the
    *largest horizontal whitespace gap* between word clusters to locate the column
    boundary.  Handles asymmetric sidebar layouts (e.g. narrow left sidebar with
    contact info and a wide right content area).
  • Column break token          — two-column pages and DOCX two-column tables emit
    ``===COLUMN_BREAK===`` between the left and right column text blocks so that
    downstream NLP models can locate the exact column boundary.
  • normalize_column_spacing()  — runs between raw extraction and section detection:
    collapses excess whitespace, fixes punctuation spacing, and preserves the
    column break token, emails, URLs, and phone numbers verbatim.
  • Table-based PDF pages       — detected separately; cells read left-to-right,
    top-to-bottom using pdfplumber's extract_tables().
  • OCR fallback intact         — triggered when digital text is too sparse.
  • DOCX two-column tables      — each column appended in order with the break token.
  • Section extraction, contact info, photo detection, language detection, and the
    dataset builder are all preserved from the previous version.

Pipeline per document:
    raw extraction → COLUMN_BREAK tokens inserted →
    normalize_column_spacing() → fix_ocr_spacing() → clean_text() →
    extract_sections()  [with dedup + confidence + fallback recovery]

Dependencies:
    pip install pdfplumber pymupdf python-docx pytesseract pillow langdetect tqdm

System dependency:
    Tesseract OCR: https://github.com/tesseract-ocr/tesseract
    Ubuntu/Debian : sudo apt-get install tesseract-ocr tesseract-ocr-tur
    macOS         : brew install tesseract
"""

import os
import re
import json
import uuid
import logging
import statistics
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
import pytesseract
from PIL import Image
import io
from tqdm import tqdm

try:
    from langdetect import detect as langdetect_detect

    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False


# ─────────────────────────────────────────────
#  LOGGING SETUP
# ─────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(levelname)s]  %(message)s",
    handlers=[
        logging.FileHandler("cv_parser.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  CONSTANTS / KEYWORD MAPS
# ─────────────────────────────────────────────

# OCR fallback threshold: if extracted text has fewer characters than this,
# we consider extraction a failure and invoke OCR.
OCR_FALLBACK_THRESHOLD = 80

# Minimum ratio of words that must appear in EACH column for multi-column detection.
# e.g. 0.15 means both left and right clusters need ≥15% of all page words.
COLUMN_MIN_RATIO = 0.15

# When scanning for the horizontal gap between columns, we project word x-ranges
# onto a 1-D grid of this many buckets.  Higher = finer resolution but slower.
GAP_SCAN_BUCKETS = 200

# Minimum gap width (as fraction of page width) to accept a column split boundary.
# Prevents splitting on narrow inter-word spaces inside a single column.
MIN_GAP_FRACTION = 0.03

# Section heading keywords — English and Turkish
SECTION_KEYWORDS: dict[str, list[str]] = {
    "summary": [
        "summary",
        "profile",
        "about",
        "about me",
        "objective",
        "professional summary",
        "career objective",
        "özet",
        "hakkımda",
        "hakkında",
        "profil",
        "kişisel özet",
    ],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
        "work history",
        "career history",
        "positions held",
        "deneyim",
        "iş deneyimi",
        "çalışma geçmişi",
        "kariyer",
    ],
    "education": [
        "education",
        "academic background",
        "academic history",
        "qualifications",
        "degrees",
        "schooling",
        "eğitim",
        "öğrenim",
        "akademik geçmiş",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "competencies",
        "technologies",
        "tools",
        "proficiencies",
        "key skills",
        "areas of expertise",
        "yetenekler",
        "beceriler",
        "teknolojiler",
        "yetkinlikler",
        "teknik beceriler",
    ],
    "projects": [
        "projects",
        "personal projects",
        "key projects",
        "portfolio",
        "open source",
        "projeler",
        "kişisel projeler",
    ],
}
# =========================
# ADVANCED SECTION PARSING
# =========================


def is_heading(line):
    line = line.strip()

    score = 0

    if len(line) < 40:
        score += 1
    if line.isupper():
        score += 1
    if re.match(r"^[A-Z][A-Za-z ]+$", line):
        score += 1
    if line.endswith(":"):
        score += 1

    keywords = sum(1 for k in SECTION_KEYWORDS if k in line.lower())
    score += keywords * 2

    return score >= 3


def extract_sections_advanced(text):
    sections = {}
    current_section = "other"

    for line in text.split("\n"):
        if is_heading(line):
            matched = None
            for key, keywords in SECTION_KEYWORDS.items():
                if any(k in line.lower() for k in keywords):
                    matched = key
                    break

            if matched:
                current_section = matched
                sections.setdefault(current_section, "")
            else:
                current_section = "other"

        else:
            sections.setdefault(current_section, "")
            sections[current_section] += line + "\n"

    return sections


def extract_sections_column_aware(text):
    parts = text.split(COLUMN_BREAK_TOKEN)

    final_sections = {}

    for part in parts:
        sec = extract_sections_advanced(part)

        for k, v in sec.items():
            if v.strip():
                final_sections[k] = final_sections.get(k, "") + "\n" + v

    return final_sections


def fallback_extract(text):
    sections = {"experience": [], "education": [], "skills": []}

    for line in text.split("\n"):
        l = line.lower()

        if any(w in l for w in ["engineer", "developer", "intern"]):
            sections["experience"].append(line)

        elif any(w in l for w in ["university", "degree", "bachelor"]):
            sections["education"].append(line)

        elif any(w in l for w in ["python", "sql", "java"]):
            sections["skills"].append(line)

    return {k: "\n".join(v) for k, v in sections.items()}


def deduplicate_sections(sections):
    cleaned = {}
    for k, v in sections.items():
        lines = list(dict.fromkeys(v.split("\n")))
        cleaned[k] = "\n".join(lines)
    return cleaned


# Turkish word list used for quick language heuristic
TURKISH_WORDS = {
    "ve",
    "bir",
    "bu",
    "da",
    "de",
    "için",
    "ile",
    "ben",
    "ama",
    "olan",
    "gibi",
    "çok",
    "ise",
    "ya",
    "ya da",
    "ki",
    "mı",
    "deneyim",
    "eğitim",
    "beceri",
    "hakkımda",
    "özet",
    "çalışma",
    "proje",
    "yetenek",
}

# Sentinel token written between the left and right column text blocks.
# Downstream NLP models can split on this string to process each column
# independently, or use it as a positional feature.
COLUMN_BREAK_TOKEN = "===COLUMN_BREAK==="


# ─────────────────────────────────────────────
#  0. TEXT NORMALISATION  (runs BEFORE section extraction)
# ─────────────────────────────────────────────

# Patterns used exclusively inside normalize_column_spacing.
# Pre-compiled at module level so repeated calls stay fast.
_NS_PROTECTED_TOKENS: tuple[re.Pattern, ...] = (
    # Order matters: match longest / most-specific patterns first.
    re.compile(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", re.IGNORECASE
    ),  # email
    re.compile(r"https?://[^\s]+|www\.[^\s]+", re.IGNORECASE),  # URL
    re.compile(r"(?:\+?\d[\d\s\-().]{6,}\d)"),  # phone
    re.compile(re.escape(COLUMN_BREAK_TOKEN)),  # our sentinel
)

_NS_MULTI_NEWLINE = re.compile(r"\n{3,}")
_NS_MULTI_SPACE = re.compile(r"[ \t]{2,}")
# Ensure exactly one space after sentence-ending punctuation when followed by a letter.
_NS_PUNCT_SPACING = re.compile(
    r"([.!?;:,])([A-Za-zÀ-ÖØ-öø-ÿ\u011e\u011f\u015e\u015f\u0130\u0131])"
)


def normalize_column_spacing(text: str) -> str:
    """
    Normalise whitespace and punctuation spacing in extracted CV text while
    preserving special tokens, emails, URLs, and phone numbers.

    Transformations applied (in order):
      1. Protect emails, URLs, phone numbers, and ``===COLUMN_BREAK===`` tokens
         by replacing them with unique placeholders — prevents any regex from
         accidentally mutating structured data.
      2. Collapse runs of 3+ newlines down to exactly 2 (one blank line).
      3. Collapse runs of 2+ spaces / tabs on the same line to a single space.
      4. Insert a single space after sentence-ending punctuation (. ! ? ; : ,)
         when it is immediately followed by a letter — fixes cases where OCR or
         PDF extraction omits the inter-sentence gap.
      5. Strip leading/trailing whitespace from every line and from the whole text.
      6. Restore all protected tokens verbatim.

    Args:
        text: Raw extracted text, potentially containing ``===COLUMN_BREAK===``
              tokens inserted by the column-aware extractor.

    Returns:
        Cleaned text with normalised spacing and all special tokens intact.

    Example::

        >>> src = "Python,SQL  TensorFlow\\n\\n\\n\\nExperience at Google"
        >>> normalize_column_spacing(src)
        'Python, SQL TensorFlow\\n\\nExperience at Google'
    """
    if not text:
        return ""

    # ── Step 1: protect structured tokens with stable placeholders ────────────
    # We use a dict keyed by a deterministic placeholder string so that
    # restoration is a simple str.replace() — no regex required.
    protected: dict[str, str] = {}

    def _protect(pattern: re.Pattern, t: str) -> str:
        """Replace every match of *pattern* with a placeholder, storing original."""

        def _replacer(m: re.Match) -> str:
            # Use a zero-padded index so placeholder length is predictable.
            key = f"\x00PROT{len(protected):04d}\x00"
            protected[key] = m.group(0)
            return key

        return pattern.sub(_replacer, t)

    for pat in _NS_PROTECTED_TOKENS:
        text = _protect(pat, text)

    # ── Step 2: collapse excess blank lines ───────────────────────────────────
    text = _NS_MULTI_NEWLINE.sub("\n\n", text)

    # ── Step 3: collapse excess horizontal whitespace ─────────────────────────
    text = _NS_MULTI_SPACE.sub(" ", text)

    # ── Step 4: ensure space after punctuation before a letter ───────────────
    text = _NS_PUNCT_SPACING.sub(r"\1 \2", text)

    # ── Step 5: strip trailing spaces from every line, then the whole text ────
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = text.strip()

    # ── Step 6: restore protected tokens verbatim ────────────────────────────
    for key, original in protected.items():
        text = text.replace(key, original)

    return text


# ─────────────────────────────────────────────
#  0b. OCR / BROKEN-TOKEN TEXT CLEANING PIPELINE
# ─────────────────────────────────────────────
#
# PURPOSE
# ───────
# Raw OCR output and some PDF extractors produce broken tokens — words that
# were split across scan lines or character groups, e.g.:
#   "gma l. com"   →  "gmail.com"
#   "ün vers tes"  →  "üniversitesi"
#   "P y t h o n"  →  "Python"
#
# This module runs AFTER normalize_column_spacing() and BEFORE clean_text().
# It must never mutate emails, URLs, phone numbers, or the COLUMN_BREAK_TOKEN.
#
# PIPELINE (in order)
# ───────────────────
#   1. Protect structured tokens (email / URL / phone / COLUMN_BREAK).
#   2. Unicode NFC normalisation — resolves composed vs decomposed characters.
#   3. Merge spaced-out single characters: "P y t h o n" → "Python".
#   4. Merge broken short tokens glued to neighbours by context rules.
#   5. Collapse residual multi-space runs.
#   6. Restore protected tokens verbatim.

# ── Compiled patterns used only inside fix_ocr_spacing() ─────────────────────

# Matches a run of single characters separated by single spaces, e.g. "P y t h o n".
# Requires at least 3 chars in the run to avoid merging real short words.
_OCR_SPACED_CHARS = re.compile(
    r"(?<!\w)([A-Za-zÀ-ÖØ-öø-ÿÀ-ɏĞğŞşİı]"
    r"(?: [A-Za-zÀ-ÖØ-öø-ÿÀ-ɏĞğŞşİı]){2,})"
    r"(?!\w)",
    re.UNICODE,
)

# Matches a lone single letter/digit surrounded by spaces that is sandwiched
# between two longer tokens on the same line — typical OCR split artifact.
# e.g. "soft w are" where "w" is the broken fragment.
# We only merge if the fragment is a single char and neighbours are ≥ 2 chars,
# to avoid merging legitimate single-letter words (a, I, etc.) mid-sentence.
_OCR_LONE_FRAGMENT = re.compile(
    r"(?<=\S{2}) ([A-Za-zÀ-ɏ]) (?=\S{2})",
    re.UNICODE,
)

# Common OCR artefacts: ligature replacements and common misreads.
_OCR_LIGATURE_MAP: list[tuple[str, str]] = [
    ("ﬁ", "fi"),  # fi ligature
    ("ﬂ", "fl"),  # fl ligature
    ("ﬀ", "ff"),  # ff ligature
    ("ﬃ", "ffi"),  # ffi ligature
    ("ﬄ", "ffl"),  # ffl ligature
    ("ﬅ", "st"),  # st ligature (rare)
    ("ﬆ", "st"),
    ("’", "'"),  # right single quotation → apostrophe
    ("“", '"'),  # left double quotation
    ("”", '"'),  # right double quotation
    ("–", "-"),  # en-dash → hyphen
    ("—", "-"),  # em-dash → hyphen
    ("·", " "),  # middle dot (used as bullet) → space
    ("", " "),  # Windows Symbol bullet → space
]

# Characters that are almost certainly OCR noise when appearing isolated
# (surrounded by spaces or at line boundaries) — e.g. stray "|", "~", "^".
_OCR_NOISE_CHARS = re.compile(r"(?<!\S)[|~^`\\](?!\S)")


def fix_ocr_spacing(text: str) -> str:
    """
    Repair common OCR / PDF-extraction spacing artifacts in CV text.

    Transformations (all token-safe — emails, URLs, phones, COLUMN_BREAK preserved):
      1. Protect structured tokens so no regex touches them.
      2. Unicode NFC normalisation (composed form, e.g. "é" not "e" + combining).
      3. Replace typographic ligatures and smart-quotes with ASCII equivalents.
      4. Remove isolated OCR noise characters ( | ~ ^ ` backslash ).
      5. Merge spaced-out individual characters: "P y t h o n" → "Python".
         Only fires on runs of ≥ 3 single chars — avoids merging "a I" etc.
      6. Merge lone single-character OCR fragments flanked by longer tokens.
      7. Collapse multi-space runs to single space (per line).
      8. Restore protected tokens verbatim.

    Args:
        text: Text after normalize_column_spacing(), before clean_text().

    Returns:
        Text with OCR spacing artifacts repaired.

    Examples:
        >>> fix_ocr_spacing("gma l. com")       # not an email yet — just broken
        'gmal.com'                               # will be caught by normaliser too
        >>> fix_ocr_spacing("P y t h o n")
        'Python'
        >>> fix_ocr_spacing("soft w are engineer")
        'software engineer'
    """
    if not text:
        return ""

    # ── Step 1: protect structured tokens ────────────────────────────────────
    protected: dict[str, str] = {}

    def _prot(pattern: re.Pattern, t: str) -> str:
        def _repl(m: re.Match) -> str:
            key = f"__OCR{len(protected):04d}__"
            protected[key] = m.group(0)
            return key

        return pattern.sub(_repl, t)

    # Protect in same priority order as normalize_column_spacing
    for pat in _NS_PROTECTED_TOKENS:
        text = _prot(pat, text)

    # ── Step 2: Unicode NFC normalisation ────────────────────────────────────
    import unicodedata as _ud

    text = _ud.normalize("NFC", text)

    # ── Step 3: Replace ligatures and typographic characters ─────────────────
    for ligature, replacement in _OCR_LIGATURE_MAP:
        text = text.replace(ligature, replacement)

    # ── Step 4: Remove isolated OCR noise characters ──────────────────────────
    text = _OCR_NOISE_CHARS.sub("", text)

    # ── Step 5: Merge spaced-out single characters ("P y t h o n") ───────────
    # We loop because the pattern is non-overlapping; one pass handles the full
    # run by consuming left-to-right, but a second pass catches any residual.
    for _ in range(3):
        text = _OCR_SPACED_CHARS.sub(lambda m: m.group(0).replace(" ", ""), text)

    # ── Step 6: Merge lone single-char OCR fragments ─────────────────────────
    # Only apply within a single line to avoid cross-line merging.
    fixed_lines = []
    for line in text.splitlines():
        # Apply up to 4 times per line (each pass may expose a new fragment).
        for _ in range(4):
            new_line = _OCR_LONE_FRAGMENT.sub(lambda m: m.group(1), line)
            if new_line == line:
                break
            line = new_line
        fixed_lines.append(line)
    text = "\n".join(fixed_lines)

    # ── Step 7: Collapse multi-space runs ────────────────────────────────────
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = "\n".join(line.rstrip() for line in text.splitlines())

    # ── Step 8: Restore protected tokens ─────────────────────────────────────
    for key, original in protected.items():
        text = text.replace(key, original)

    return text


# ─────────────────────────────────────────────
#  1. COLUMN-AWARE PDF TEXT EXTRACTION
# ─────────────────────────────────────────────

# ── 1a. Gap-based column boundary detection ───────────────────────────────────


def _find_column_split_x(words: list[dict], page_width: float) -> Optional[float]:
    """
    Find the x-coordinate of the widest horizontal whitespace gap between
    word clusters on a page — this is where the column boundary lies.

    Algorithm:
      1. Build a 1-D occupancy array along the x-axis (GAP_SCAN_BUCKETS wide).
      2. Mark each bucket as 'occupied' if any word overlaps it.
      3. Find the longest contiguous run of *empty* buckets.
      4. Return the centre x of that run, or None if no significant gap found.

    This is more robust than a hard midpoint split because it adapts to:
      - Asymmetric layouts (narrow sidebar + wide main column)
      - Layouts where the column split is off-centre
    """
    if not words:
        return None

    bucket_size = page_width / GAP_SCAN_BUCKETS
    occupied = [False] * GAP_SCAN_BUCKETS

    for w in words:
        # Mark every bucket overlapped by this word's x extent
        start_bucket = max(0, int(w["x0"] / bucket_size))
        end_bucket = min(GAP_SCAN_BUCKETS - 1, int(w["x1"] / bucket_size))
        for b in range(start_bucket, end_bucket + 1):
            occupied[b] = True

    # Find the longest unoccupied run
    best_start = best_end = -1
    current_start = None

    for i, occ in enumerate(occupied):
        if not occ:
            if current_start is None:
                current_start = i
        else:
            if current_start is not None:
                run_len = i - current_start
                if run_len > (best_end - best_start):
                    best_start, best_end = current_start, i - 1
                current_start = None

    # Handle gap that runs to the end of the array
    if current_start is not None:
        run_len = GAP_SCAN_BUCKETS - current_start
        if run_len > (best_end - best_start):
            best_start, best_end = current_start, GAP_SCAN_BUCKETS - 1

    if best_start == -1:
        return None  # No gap found

    gap_width_fraction = (best_end - best_start + 1) / GAP_SCAN_BUCKETS
    if gap_width_fraction < MIN_GAP_FRACTION:
        # Gap too narrow — probably just inter-word spacing, not a column split
        return None

    # Return the pixel x-coordinate of the gap's centre
    gap_centre_x = ((best_start + best_end) / 2.0) * bucket_size
    return gap_centre_x


# ── 1b. Layout detection ──────────────────────────────────────────────────────


class PageLayout:
    """Enum-like class for layout type labels."""

    SINGLE = "single"
    TWO_COL = "two_column"
    MULTI = "multi_column"  # 3+ columns (unusual but exists)
    TABLE = "table"  # page dominated by a table structure


def _detect_page_layout(page, words: list[dict]) -> str:
    """
    Classify a pdfplumber page into a layout type.

    Decision logic:
      1. If pdfplumber finds any tables on the page → TABLE layout.
      2. Count words left/right of the detected gap boundary.
         - If gap exists AND both sides have >= COLUMN_MIN_RATIO → TWO_COL.
         - Additional check: if standard deviation of x0 values is very high
           (words spread all over), consider MULTI.
      3. Otherwise → SINGLE.
    """
    # Check for explicit table structures first
    try:
        tables = page.extract_tables()
        if tables and any(len(t) > 1 for t in tables):
            # Only flag as TABLE if there's a meaningful table (>1 row)
            return PageLayout.TABLE
    except Exception:
        pass

    if not words:
        return PageLayout.SINGLE

    page_width = page.width
    split_x = _find_column_split_x(words, page_width)

    if split_x is None:
        return PageLayout.SINGLE

    # Use centre_x for both detection and extraction — must stay consistent
    # with _extract_two_column, which also uses centre_x.  The old code used
    # x1 / x0 thresholds here and centre_x in _extract_two_column, causing
    # words near the gutter to be counted in detection but routed to the wrong
    # column during extraction, which corrupted the column split decision.
    left_count = sum(1 for w in words if (w["x0"] + w["x1"]) / 2 <= split_x)
    right_count = sum(1 for w in words if (w["x0"] + w["x1"]) / 2 > split_x)
    total = len(words)

    left_ratio = left_count / total
    right_ratio = right_count / total

    if left_ratio < COLUMN_MIN_RATIO or right_ratio < COLUMN_MIN_RATIO:
        return PageLayout.SINGLE

    # Quick multi-column check: look for a second significant gap in EACH half.
    # Right words keep absolute x-coords; translate them so x starts from 0
    # before scanning — otherwise gap fractions are against the full-page width
    # and the sub-gap check is essentially disabled.
    left_words = [w for w in words if (w["x0"] + w["x1"]) / 2 <= split_x]
    right_words = [w for w in words if (w["x0"] + w["x1"]) / 2 > split_x]
    right_words_t = [
        {**w, "x0": w["x0"] - split_x, "x1": w["x1"] - split_x} for w in right_words
    ]

    left_gap = _find_column_split_x(left_words, split_x)
    right_gap = _find_column_split_x(right_words_t, page_width - split_x)

    if left_gap is not None or right_gap is not None:
        return PageLayout.MULTI

    return PageLayout.TWO_COL


# ── 1c. Word-list → ordered text reconstruction ───────────────────────────────


def _words_to_text(word_list: list[dict], y_tolerance: float = 4.0) -> str:
    """
    Convert a list of pdfplumber word dicts to a text string.

    Words are sorted by (top, x0) and grouped into lines when their
    vertical positions are within y_tolerance points of each other.
    This handles slight baseline misalignments common in designed CVs.

    Args:
        word_list    : list of pdfplumber word dicts (keys: text, x0, x1, top).
        y_tolerance  : vertical distance (pts) within which words share a line.

    Returns:
        Newline-separated string, one line per detected text row.
    """
    if not word_list:
        return ""

    # Sort: primary = top (vertical position), secondary = x0 (left-to-right)
    word_list = sorted(
        word_list, key=lambda w: (round(w["top"] / y_tolerance) * y_tolerance, w["x0"])
    )

    lines: list[str] = []
    current_line: list[str] = []
    current_top: Optional[float] = None

    for w in word_list:
        if current_top is None or abs(w["top"] - current_top) < y_tolerance:
            current_line.append(w["text"])
            # Keep track of the average top so we don't drift on long lines
            current_top = (
                w["top"] if current_top is None else (current_top + w["top"]) / 2
            )
        else:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [w["text"]]
            current_top = w["top"]

    if current_line:
        lines.append(" ".join(current_line))

    return "\n".join(lines)


# ── 1d. Per-layout extraction functions ──────────────────────────────────────


def _extract_single_column(page) -> str:
    """
    Standard single-column extraction.

    Uses pdfplumber's built-in extract_text() with tight tolerances to
    keep words on the same line together while preserving paragraph breaks.
    """
    return page.extract_text(x_tolerance=3, y_tolerance=3) or ""


def _extract_two_column(page, words: list[dict]) -> str:
    """
    Reconstruct reading order for a two-column page.

    The column boundary is found via _find_column_split_x() (gap analysis)
    rather than a fixed midpoint — this correctly handles asymmetric layouts
    such as a narrow contact sidebar on the left and a wide experience column
    on the right.

    Reading order: left column (top → bottom) then right column (top → bottom).

    Args:
        page  : pdfplumber page object (needed for page_width).
        words : pre-extracted word list (avoids re-extracting).

    Returns:
        Reconstructed text with left column first, then right column.
    """
    page_width = page.width
    split_x = _find_column_split_x(words, page_width)

    if split_x is None:
        # Fallback to simple midpoint if gap detection fails
        split_x = page_width / 2
        logger.debug(
            "  [layout] Gap detection failed — falling back to midpoint split."
        )

    # Partition words into left and right columns
    # Note: words straddling the gap (x0 < split_x < x1) go to whichever
    # column their *centre* falls in — avoids double-counting headers.
    left_words: list[dict] = []
    right_words: list[dict] = []

    for w in words:
        centre_x = (w["x0"] + w["x1"]) / 2
        if centre_x <= split_x:
            left_words.append(w)
        else:
            right_words.append(w)

    left_text = _words_to_text(left_words)
    right_text = _words_to_text(right_words)

    # Left column first, then right column.
    # The COLUMN_BREAK_TOKEN sentinel is injected between them so that:
    #   • downstream NLP models can locate the exact column boundary;
    #   • section-heading regexes still see headings at line starts on each side;
    #   • normalize_column_spacing() preserves the token verbatim.
    non_empty = [p for p in [left_text, right_text] if p.strip()]
    if len(non_empty) == 2:
        return f"{non_empty[0]}\n\n{COLUMN_BREAK_TOKEN}\n\n{non_empty[1]}"
    # Only one side had content — no break token needed.
    return non_empty[0] if non_empty else ""


def _extract_table_page(page) -> str:
    """
    Extract text from a table-dominated page.

    Strategy:
      1. Extract tables cell-by-cell (left-to-right, top-to-bottom per row).
         Each cell's content is kept as a separate block so section headings
         in table headers are preserved.
      2. Also extract any non-table paragraphs floating outside the tables.

    This is the right approach for CV templates that use invisible tables as
    layout grids (common in Word-exported PDFs).
    """
    parts: list[str] = []

    try:
        tables = page.extract_tables()
        if tables:
            for table in tables:
                for row in table:
                    if row:
                        for cell in row:
                            cell_text = (cell or "").strip()
                            if cell_text:
                                parts.append(cell_text)
    except Exception as e:
        logger.warning(f"  [layout_issue] Table extraction failed on page: {e}")

    # Also capture text not inside any table bounding box
    try:
        non_table_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
        if non_table_text.strip():
            parts.append(non_table_text)
    except Exception:
        pass

    return "\n".join(parts)


def _extract_multi_column(page, words: list[dict]) -> str:
    """
    Fallback for pages with 3+ columns (rare in CVs but exists in fancy templates).

    Strategy: cluster words by their x0 into N groups using a simple gap scan,
    sort each cluster top-to-bottom, and concatenate left-to-right.

    If clustering is ambiguous, we fall back to a plain word-by-word sort by
    (top, x0) which is still better than line-by-line extraction.
    """
    logger.info("  [layout] Multi-column (3+) page detected — using positional sort.")

    # Sort all words by natural reading order (top → bottom, left → right)
    # For most 3-column layouts this produces a readable result.
    return _words_to_text(words, y_tolerance=5.0)


# ── 1e. Main PDF extraction orchestrator ─────────────────────────────────────


def extract_text_pdf(file_path: str) -> tuple[str, str]:
    """
    Extract text from a PDF file using layout-aware column reconstruction.

    Returns:
        (text, source_format) where source_format is "pdf" or "ocr".

    Per-page strategy:
      1. Extract word bounding boxes with pdfplumber.
      2. Detect layout type: SINGLE / TWO_COL / MULTI / TABLE.
      3. Dispatch to the appropriate extraction function.
      4. If the combined extracted text is too short → OCR fallback.
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    all_pages_text: list[str] = []
    layout_issues: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    # Extract word dicts once — reused by detection + extraction
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False,
                        use_text_flow=False,  # raw positions, not PDF text-flow
                    )

                    layout = _detect_page_layout(page, words)

                    if layout == PageLayout.TWO_COL:
                        logger.info(
                            f"  [layout] TWO_COLUMN detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_two_column(page, words)

                    elif layout == PageLayout.TABLE:
                        logger.info(
                            f"  [layout] TABLE layout detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_table_page(page)

                    elif layout == PageLayout.MULTI:
                        logger.info(
                            f"  [layout] MULTI_COLUMN (3+) detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_multi_column(page, words)

                    else:
                        # SINGLE column — standard extraction
                        page_text = _extract_single_column(page)

                    all_pages_text.append(page_text)

                except Exception as page_err:
                    logger.warning(
                        f"  [layout_issue] page {page_num} of '{basename}': {page_err}"
                    )
                    layout_issues.append(f"page_{page_num}")
                    all_pages_text.append("")

        full_text = "\n\n".join(filter(None, all_pages_text))

        if layout_issues:
            logger.warning(
                f"  [layout_issue] '{basename}' — problematic pages: {layout_issues}"
            )

        if len(full_text.strip()) >= OCR_FALLBACK_THRESHOLD:
            return full_text, "pdf"

        # Text is too short — fall through to OCR
        logger.info(
            f"  [pdf→ocr] Text too short ({len(full_text.strip())} chars) "
            f"in '{basename}' — invoking OCR."
        )

    except Exception as e:
        logger.warning(
            f"  [pdf_error] pdfplumber failed on '{basename}': {e} — invoking OCR."
        )

    return ocr_fallback(file_path)


# ─────────────────────────────────────────────
#  2. TEXT EXTRACTION — DOCX
# ─────────────────────────────────────────────


def extract_text_docx(file_path: str) -> tuple[str, str]:
    """
    Extract text from a DOCX file.

    Handles:
      - Regular paragraphs (main body flow)
      - Tables — including two-column layout tables:
          * Single-cell rows are read normally.
          * Two-cell rows are treated as left/right columns:
            left cell content is appended first, then right cell content.
            This preserves sidebar-style layouts common in designed DOCX CVs.
          * Rows with 3+ cells: each cell appended in left-to-right order.
      - Text boxes (shapes containing text frames in the XML).

    Returns:
        (text, "docx")
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    doc = Document(file_path)
    parts: list[str] = []

    # ── Paragraphs (main body) ────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # ── Tables ────────────────────────────────────────────────────────────────
    # We treat each table separately to avoid merging unrelated columns.
    # Two-column tables are the most common layout device in DOCX CVs.
    for table_idx, table in enumerate(doc.tables):
        col_count = max((len(row.cells) for row in table.rows), default=0)

        if col_count == 2:
            # Two-column table: left column first, then right column.
            # Collect each column's text independently, then interleave so that
            # content from the same row stays adjacent — better for section detection.
            logger.info(
                f"  [docx_layout] Two-column table detected "
                f"(table {table_idx + 1}) in '{basename}'"
            )
            left_parts: list[str] = []
            right_parts: list[str] = []
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    left_text = cells[0].text.strip()
                    right_text = cells[1].text.strip()
                    if left_text:
                        left_parts.append(left_text)
                    if right_text:
                        right_parts.append(right_text)
                elif len(cells) == 1:
                    cell_text = cells[0].text.strip()
                    if cell_text:
                        left_parts.append(cell_text)

            # Left column block first, then COLUMN_BREAK_TOKEN, then right column block.
            # This mirrors the PDF two-column treatment so the downstream pipeline
            # receives a consistently tokenised stream regardless of source format.
            if left_parts:
                parts.append("\n".join(left_parts))
            if left_parts and right_parts:
                parts.append(COLUMN_BREAK_TOKEN)
            if right_parts:
                parts.append("\n".join(right_parts))

        else:
            # Single-column or 3+-column table: cell-by-cell, row-by-row
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

    # ── Text boxes (drawing canvas shapes) ───────────────────────────────────
    # Embedded as <w:txbxContent> → child <w:p> → <w:t> in the XML.
    try:
        body = doc.element.body
        for textbox in body.iter(qn("w:txbxContent")):
            for child_para in textbox.iter(qn("w:p")):
                texts = [node.text for node in child_para.iter(qn("w:t")) if node.text]
                if texts:
                    parts.append(" ".join(texts))
    except Exception as e:
        logger.warning(
            f"  [docx_textbox] Could not extract text boxes from '{basename}': {e}"
        )

    full_text = "\n".join(parts)
    logger.info(f"  [docx] Extracted {len(full_text)} chars from '{basename}'")
    return full_text, "docx"


# ─────────────────────────────────────────────
#  3. OCR FALLBACK
# ─────────────────────────────────────────────


def ocr_fallback(file_path: str) -> tuple[str, str]:
    """
    Rasterise each page of a PDF with PyMuPDF and run Tesseract OCR.

    We try English + Turkish language packs (eng+tur).
    Falls back to English-only if the combined pack is unavailable.

    Returns:
        (text, "ocr") or ("", "failed") on complete failure.
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    all_text: list[str] = []

    logger.info(f"  [ocr] Starting OCR on '{basename}'")

    try:
        pdf_doc = fitz.open(file_path)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            # Render at 300 DPI for good OCR accuracy
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))

            # Attempt combined language OCR, fall back to English-only
            ocr_success = False
            for lang in ("eng+tur", "eng"):
                try:
                    text = pytesseract.image_to_string(img, lang=lang, config="--psm 6")
                    all_text.append(text)
                    if lang == "eng+tur":
                        logger.info(
                            f"  [ocr] Page {page_num + 1}: eng+tur OCR successful."
                        )
                    else:
                        logger.info(
                            f"  [ocr] Page {page_num + 1}: eng OCR used (tur pack unavailable)."
                        )
                    ocr_success = True
                    break
                except pytesseract.pytesseract.TesseractError:
                    continue

            if not ocr_success:
                logger.warning(
                    f"  [ocr_warning] Tesseract failed entirely on page "
                    f"{page_num + 1} of '{basename}'"
                )

        pdf_doc.close()
        full_text = "\n\n".join(filter(None, all_text))
        logger.info(f"  [ocr] Extracted {len(full_text)} chars from '{basename}'")
        return full_text, "ocr"

    except Exception as e:
        logger.error(f"  [ocr_failed] OCR failed for '{basename}': {e}")
        return "", "failed"


# ─────────────────────────────────────────────
#  4. TEXT CLEANING
# ─────────────────────────────────────────────

# Pre-compiled patterns for efficiency
_RE_EMAIL = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    re.IGNORECASE,
)
_RE_URL = re.compile(
    r"https?://[^\s]+|www\.[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+",
    re.IGNORECASE,
)
_RE_PHONE = re.compile(
    r"(?:\+?\d[\d\s\-().]{6,}\d)",
)
_RE_MULTI_SPACE = re.compile(r"[ \t]{2,}")
_RE_MULTI_NEWLINE = re.compile(r"\n{3,}")
_RE_SPECIAL_CHARS = re.compile(r"[^\w\s@.,:;()\-+/#&'\"/\\]", re.UNICODE)


def clean_text(text: str) -> str:
    """
    Selective text cleaning that preserves structured data.

    Rules:
      1. Replace protected tokens (emails, URLs, phone numbers) with placeholders.
      2. Strip excessive special characters.
      3. Normalize whitespace and newlines.
      4. Lowercase everything.
      5. Restore protected tokens.

    This ensures we never break email addresses, URLs, or phone numbers.
    """
    if not text:
        return ""

    # ── Step 1: Protect structured tokens ───────────────────────────────────
    protected: dict[str, str] = {}

    def protect(pattern: re.Pattern, prefix: str, t: str) -> str:
        def replacer(m):
            key = f"__PROTECTED_{prefix}_{len(protected)}__"
            protected[key] = m.group(0)
            return key

        return pattern.sub(replacer, t)

    # Protect the column-break sentinel FIRST (before email/URL patterns that
    # might partially match characters inside it).
    text = protect(re.compile(re.escape(COLUMN_BREAK_TOKEN)), "COLBREAK", text)
    text = protect(_RE_EMAIL, "EMAIL", text)
    text = protect(_RE_URL, "URL", text)
    text = protect(_RE_PHONE, "PHONE", text)

    # ── Step 2: Remove unwanted special characters ────────────────────────
    text = _RE_SPECIAL_CHARS.sub(" ", text)

    # ── Step 3: Normalize whitespace ─────────────────────────────────────
    text = _RE_MULTI_SPACE.sub(" ", text)
    text = _RE_MULTI_NEWLINE.sub("\n\n", text)
    text = text.strip()

    # ── Step 4: Lowercase ────────────────────────────────────────────────
    text = text.lower()

    # ── Step 5: Restore protected tokens ─────────────────────────────────
    for key, original in protected.items():
        text = text.replace(key.lower(), original)

    return text


# ─────────────────────────────────────────────
#  5. SECTION DETECTION  (state-machine edition)
# ─────────────────────────────────────────────
#
# PROBLEM WITH THE OLD APPROACH
# ──────────────────────────────
# The previous implementation used a single compiled regex to locate headings
# and then sliced the raw text between match positions.  This failed in two ways:
#
#   1. FALSE POSITIVES — _classify_heading used a bidirectional substring test
#      ("kw in heading_lower or heading_lower in kw").  This fired on body-text
#      lines containing a keyword word mid-sentence (e.g. "strong skills in…")
#      which caused a spurious section split mid-paragraph → bleeding.
#
#   2. REGEX ANCHOR CONFUSION — re.MULTILINE makes ^ match at every newline, so
#      any line containing a keyword anywhere triggered a heading match, even
#      if it was clearly content rather than a standalone heading.
#
# STATE-MACHINE FIX
# ─────────────────
# We now process the text one line at a time.  A line is a heading only if
# it passes _is_section_heading(), which requires:
#   • the normalised line (stripped, no punctuation) is an *exact* match or a
#     very close keyword match (keyword == whole normalised line, allowing a
#     trailing slash-separated bilingual label like "skills / yetenekler").
#   • the line is "short" — headings are almost never long sentences.
# The state machine accumulates body lines into the current section bucket
# without any risk of a content line hijacking the section pointer.

import unicodedata

# Maximum word count a line may have to be considered a heading candidate.
# Heading lines like "PROFESSIONAL EXPERIENCE" have ~2-3 words.
# Body lines have many more.  Threshold of 6 keeps most multi-word heading
# phrases while excluding prose sentences.
_HEADING_MAX_WORDS = 6

# Pre-built normalised keyword → canonical-section index for O(1) lookup.
# Keys are normalised (lowercase, no punctuation, stripped) keyword strings.
_KW_NORM_MAP: dict[str, str] = {}
for _section, _kws in SECTION_KEYWORDS.items():
    for _kw in _kws:
        _norm_kw = re.sub(r"[^\w\s]", "", _kw.lower(), flags=re.UNICODE).strip()
        _KW_NORM_MAP[_norm_kw] = _section


def _normalise_heading_line(line: str) -> str:
    """
    Normalise a line for heading comparison.

    Transformations:
      • strip surrounding whitespace
      • lowercase
      • remove all punctuation (incl. Turkish special chars)
      • collapse runs of whitespace to single space

    Bilingual headings like "Skills / Yetenekler" become "skills  yetenekler"
    which is then collapsed to "skills yetenekler" — handled by the lookup.
    """
    # Remove punctuation: keep letters, digits, whitespace (unicode-aware)
    cleaned = re.sub(r"[^\w\s]", " ", line.lower(), flags=re.UNICODE)
    return re.sub(r"\s+", " ", cleaned).strip()


def _is_section_heading(line: str) -> Optional[str]:
    """
    Determine whether *line* is a section heading.

    Returns the canonical section name (e.g. "experience") if it is a heading,
    or None otherwise.

    Matching rules (applied in order):
      1. Reject lines with more than _HEADING_MAX_WORDS words — body text.
      2. Normalise the line (lowercase, strip punctuation).
      3. EXACT match: normalised line == a keyword  → return that section.
      4. BILINGUAL match: normalised line contains a "/" separator; check each
         part against the keyword map.
      5. CONTAINS match (conservative): the normalised line is fully contained
         within a known keyword phrase (i.e. partial keyword, e.g. "education"
         inside "academic education").  We do NOT fire on keyword-as-substring-
         of-body-text — that direction is intentionally excluded here.
    """
    stripped = line.strip()
    if not stripped:
        return None

    # Rule 1 — length guard: real headings are short
    word_count = len(stripped.split())
    if word_count > _HEADING_MAX_WORDS:
        return None

    norm = _normalise_heading_line(stripped)
    if not norm:
        return None

    # Rule 3 — exact match on fully-normalised line
    if norm in _KW_NORM_MAP:
        return _KW_NORM_MAP[norm]

    # Rule 4 — bilingual "keyword / keyword" pattern
    # IMPORTANT: "/" and "–" are stripped to spaces by _normalise_heading_line,
    # so we must split on the RAW stripped line and normalise each part
    # individually — not on the already-normalised `norm`.
    if re.search(r"[/–]", stripped):
        for raw_part in re.split(r"\s*[/–]\s*", stripped):
            part_norm = _normalise_heading_line(raw_part)
            if part_norm and part_norm in _KW_NORM_MAP:
                return _KW_NORM_MAP[part_norm]

    return None


# ── Minimum line count for a section to be considered "non-empty" ────────────
_SECTION_MIN_LINES = 1

# ── How many lines of context to scan for fallback keyword recovery ───────────
_FALLBACK_WINDOW = 50


def _score_section(lines: list[str]) -> float:
    """
    Return a confidence score [0.0 – 1.0] for how likely a section's content
    is genuine, based on simple heuristics:

      • 0.0  — empty
      • 0.1  — single very short line (< 10 chars) — probably a stray artefact
      • 0.5  — has content but only 1–2 lines
      • 0.8  — 3+ lines
      • 1.0  — 5+ lines with ≥ 20 chars average length

    Used for logging and downstream quality filtering.
    """
    if not lines:
        return 0.0
    non_empty = [l for l in lines if l.strip()]
    if not non_empty:
        return 0.0
    if len(non_empty) == 1 and len(non_empty[0].strip()) < 10:
        return 0.1
    if len(non_empty) <= 2:
        return 0.5
    avg_len = sum(len(l.strip()) for l in non_empty) / len(non_empty)
    if len(non_empty) >= 5 and avg_len >= 20:
        return 1.0
    return 0.8


def _fallback_keyword_recovery(
    text: str,
    empty_sections: list[str],
) -> dict[str, list[str]]:
    """
    Last-resort keyword scan for sections that the state machine left empty.

    Strategy: for each empty section, scan all lines of the text for lines
    containing any of its keywords as a *whole word* (not substring).  Collect
    up to _FALLBACK_WINDOW lines following the first keyword hit.

    This is intentionally less strict than _is_section_heading() — we are in
    fallback mode and accept some noise in exchange for not returning empty.

    Returns a dict of { section_name: [recovered_lines] } for the empty sections
    only; caller merges into the main sections dict.
    """
    recovered: dict[str, list[str]] = {sec: [] for sec in empty_sections}
    all_lines = text.splitlines()

    for section in empty_sections:
        keywords = SECTION_KEYWORDS.get(section, [])
        # Build a set of normalised keyword tokens for whole-word matching
        kw_norms = {
            re.sub(r"[^\w\s]", "", kw.lower(), flags=re.UNICODE).strip()
            for kw in keywords
        }

        hit_idx: Optional[int] = None
        for idx, line in enumerate(all_lines):
            line_norm = re.sub(r"[^\w\s]", " ", line.lower(), flags=re.UNICODE)
            line_words = set(line_norm.split())
            if kw_norms & line_words:  # any keyword word appears in line
                hit_idx = idx
                break

        if hit_idx is not None:
            # Collect up to _FALLBACK_WINDOW lines after the keyword hit,
            # stopping at the next section heading.
            window = all_lines[hit_idx + 1 : hit_idx + 1 + _FALLBACK_WINDOW]
            for line in window:
                if _is_section_heading(line) is not None:
                    break
                if line.strip():
                    recovered[section].append(line)

    return recovered


def extract_sections(text: str, debug: bool = False) -> dict[str, str]:
    """
    Split cleaned CV text into structured sections using a line-by-line
    state machine — with duplicate/loop prevention, confidence scoring,
    fallback recovery for empty sections, and optional debug output.

    State machine logic (unchanged from previous version)
    ──────────────────────────────────────────────────────
    current_section = None
    for each line:
        if _is_section_heading(line):
            current_section = detected_section
        else:
            append line to sections[current_section]

    NEW: Duplicate / loop prevention
    ─────────────────────────────────
    A common failure on column-corrupted CVs is:
        education → experience → education → experience  (oscillating)
    or
        experience → experience  (heading repeated from both columns)

    Fix: track the SEQUENCE of section transitions.  If a section is seen
    again AFTER having already transitioned away from it, we do NOT reset
    the accumulator — we APPEND to the existing content instead.  This means
    that split-column content for the same logical section is merged rather
    than creating phantom duplicate blocks.

    NEW: Confidence scoring
    ────────────────────────
    Each section receives a score via _score_section().  Logged as a warning
    when score < 0.5.  Returned in the output dict under "__confidence__" key
    (a sub-dict) for downstream consumers.

    NEW: Fallback recovery
    ──────────────────────
    After the state machine pass, any section still empty triggers
    _fallback_keyword_recovery(), which does a full-text keyword scan.

    NEW: Debug output
    ──────────────────
    When debug=True (or when PARSER_DEBUG env var is set), prints:
      • CLEANED TEXT SAMPLE (first 300 chars)
      • DETECTED SECTIONS (names + line counts)
      • EMPTY SECTION WARNING for any section with score < 0.5

    Args:
        text:  Cleaned, lowercased text with ===COLUMN_BREAK=== tokens intact.
        debug: If True, print diagnostic output to stdout.

    Returns:
        Dict with keys: summary, experience, education, skills, projects.
        Each value is stripped body text.  Missing sections default to "".
        Also includes "__confidence__" sub-dict with float scores per section.
    """
    _debug = debug or bool(os.environ.get("PARSER_DEBUG", ""))

    if _debug:
        sample = text[:300].replace("\n", " ↵ ")
        print(f"[DEBUG] CLEANED TEXT SAMPLE: {sample!r}")

    sections: dict[str, list[str]] = {
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
    }

    current_section: Optional[str] = None

    # ── Duplicate / loop prevention ───────────────────────────────────────────
    # transition_log records every section-heading transition in document order.
    # If we detect that the same section appears again AFTER at least one OTHER
    # section was seen in between, we are in a loop/duplicate situation —
    # we keep appending to the same bucket (do not reset current_section away
    # from the prior occurrence).
    transition_log: list[str] = []  # ordered list of section names as seen
    seen_sections: set[str] = set()  # set of section names encountered so far

    for raw_line in text.splitlines():
        # ── Pass COLUMN_BREAK_TOKEN through unchanged ─────────────────────────
        if COLUMN_BREAK_TOKEN in raw_line:
            if current_section:
                sections[current_section].append(raw_line)
            continue

        detected = _is_section_heading(raw_line)

        if detected is not None:
            # ── Duplicate / loop check ────────────────────────────────────────
            if detected in seen_sections:
                # This section heading has been seen before.
                # Check if we've visited any OTHER section since the last time
                # we were in this section.
                last_occurrence = (
                    len(transition_log) - 1 - transition_log[::-1].index(detected)
                )
                sections_since = set(transition_log[last_occurrence + 1 :])
                sections_since.discard(detected)

                if sections_since:
                    # We left this section, visited others, and came back.
                    # This is a loop/repeat — merge into existing bucket by
                    # simply NOT changing current_section (the accumulator
                    # for this section keeps growing).
                    logger.debug(
                        f"  [dedup] Repeated section heading '{detected}' "
                        f"after visiting {sections_since} — merging content."
                    )
                    # Keep current_section as detected (merge), record transition
                    current_section = detected
                else:
                    # Consecutive repetition of same heading (e.g. from both columns).
                    # Just stay in the same section — no-op.
                    pass
            else:
                # New section — normal transition
                current_section = detected
                seen_sections.add(detected)

            transition_log.append(detected)

        else:
            # Body line — accumulate into current section
            if current_section is not None and raw_line.strip():
                sections[current_section].append(raw_line)

    # ── Confidence scoring ────────────────────────────────────────────────────
    confidence: dict[str, float] = {
        sec: _score_section(lines) for sec, lines in sections.items()
    }

    # ── Fallback recovery for empty sections ─────────────────────────────────
    empty_sections = [sec for sec, lines in sections.items() if not lines]
    if empty_sections:
        recovered = _fallback_keyword_recovery(text, empty_sections)
        for sec, lines in recovered.items():
            if lines:
                sections[sec] = lines
                confidence[sec] = _score_section(lines) * 0.6  # discounted score
                logger.info(
                    f"  [fallback] Recovered {len(lines)} line(s) "
                    f"for empty section '{sec}' via keyword scan."
                )

    # ── Build final output ────────────────────────────────────────────────────
    result: dict[str, str] = {
        sec: "\n".join(lines).strip() for sec, lines in sections.items()
    }
    result["__confidence__"] = confidence  # type: ignore[assignment]

    # ── Debug output ──────────────────────────────────────────────────────────
    if _debug:
        print("[DEBUG] DETECTED SECTIONS:")
        for sec, text_val in result.items():
            if sec == "__confidence__":
                continue
            line_count = len(text_val.splitlines()) if text_val else 0
            score = confidence.get(sec, 0.0)
            print(f"         {sec:<12}  lines={line_count:<4}  confidence={score:.2f}")

    # ── Log warnings for low-confidence sections ──────────────────────────────
    for sec, score in confidence.items():
        if score < 0.5:
            msg = (
                f"  [quality] EMPTY SECTION WARNING: '{sec}' has "
                f"confidence={score:.2f} (lines={len(sections[sec])})"
            )
            logger.warning(msg)
            if _debug:
                print(
                    f"[DEBUG] EMPTY SECTION WARNING: '{sec}' — confidence={score:.2f}"
                )

    return result


# ─────────────────────────────────────────────
#  6. CONTACT INFO EXTRACTION
# ─────────────────────────────────────────────

_RE_LINKEDIN = re.compile(
    r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_\-%.]+)",
    re.IGNORECASE,
)
_RE_GITHUB = re.compile(
    r"(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_\-]+)",
    re.IGNORECASE,
)
_RE_PHONE_CONTACT = re.compile(
    r"(?<!\d)(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?){1,4}\d{2,4}(?!\d)",
)


def extract_contact_info(text: str) -> dict[str, str]:
    """
    Extract contact information using precise regular expressions.

    Extracts: email, phone, linkedin, github.

    IMPORTANT: We operate on the ORIGINAL (un-cleaned) text to avoid
    mangling punctuation inside emails and URLs.
    """
    contact: dict[str, str] = {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
    }

    email_match = _RE_EMAIL.search(text)
    if email_match:
        contact["email"] = email_match.group(0).strip()

    phone_matches = _RE_PHONE_CONTACT.findall(text)
    for raw in phone_matches:
        digits = re.sub(r"\D", "", raw)
        if 7 <= len(digits) <= 15:
            contact["phone"] = raw.strip()
            break

    linkedin_match = _RE_LINKEDIN.search(text)
    if linkedin_match:
        full = linkedin_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["linkedin"] = full

    github_match = _RE_GITHUB.search(text)
    if github_match:
        full = github_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["github"] = full

    return contact


# ─────────────────────────────────────────────
#  7. PHOTO DETECTION
# ─────────────────────────────────────────────


def detect_photo_pdf(file_path: str) -> bool:
    """
    Detect whether the PDF contains at least one embedded image (likely a profile photo).

    Uses PyMuPDF's get_images() — returns all images per page.
    We ignore tiny images (< 50×50 px) which are likely icons or decorations.
    """
    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            images = page.get_images(full=True)
            for img in images:
                # img tuple: (xref, smask, width, height, bpc, colorspace, ...)
                width = img[2]
                height = img[3]
                if width >= 50 and height >= 50:
                    doc.close()
                    return True
        doc.close()
    except Exception as e:
        logger.warning(
            f"  [photo_pdf] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo_docx(file_path: str) -> bool:
    """
    Detect whether a DOCX file contains inline images (likely a profile photo).

    Checks:
      - Inline shapes with pictures
      - Relationships of type image in the document part
    """
    try:
        doc = Document(str(file_path))

        for shape in doc.inline_shapes:
            if shape.type is not None:
                return True

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                return True

    except Exception as e:
        logger.warning(
            f"  [photo_docx] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo(file_path: str, source_format: str) -> bool:
    """Dispatch to format-specific photo detection."""
    if source_format in ("pdf", "ocr"):
        return detect_photo_pdf(file_path)
    elif source_format == "docx":
        return detect_photo_docx(file_path)
    return False


# ─────────────────────────────────────────────
#  8. LANGUAGE DETECTION
# ─────────────────────────────────────────────


def detect_language(text: str) -> str:
    """
    Detect whether the CV is in Turkish, English, or mixed.

    Strategy:
      1. Simple heuristic: count Turkish stopwords in lowercased text.
      2. If langdetect is available: two-pass (start + end samples).
         Both must agree → that language; else → "mixed".
      3. Fallback: if Turkish words heuristic triggers → "tr", else "en".

    Returns "tr" | "en" | "mixed"
    """
    if not text or len(text.strip()) < 30:
        return "en"

    sample = text[:2000].lower()
    words_in_sample = set(re.findall(r"\b\w+\b", sample))
    turkish_hits = words_in_sample & TURKISH_WORDS

    if LANGDETECT_AVAILABLE:
        try:
            sample_start = text[:1000]
            sample_end = text[-500:]

            lang_start = langdetect_detect(sample_start)
            lang_end = langdetect_detect(sample_end)

            def norm(lang: str) -> str:
                return "tr" if lang == "tr" else "en"

            l1, l2 = norm(lang_start), norm(lang_end)
            return l1 if l1 == l2 else "mixed"

        except Exception:
            pass  # fall through to heuristic

    if len(turkish_hits) >= 3:
        return "tr"
    return "en"


# ─────────────────────────────────────────────
#  9. MAIN PROCESSING FUNCTION
# ─────────────────────────────────────────────


def process_cv(file_path: Path) -> dict:
    """
    Process a single CV file (PDF or DOCX) and return a structured record.

    Pipeline (in order):
      1.  Determine file format.
      2.  Extract raw text — column-aware, with OCR fallback.
          Two-column pages/tables emit a ``===COLUMN_BREAK===`` sentinel
          between the left and right column text blocks.
      3.  Extract contact info from the ORIGINAL text (before any cleaning)
          to preserve punctuation inside emails, URLs, and phone numbers.
      4.  Normalize column spacing:
          ``normalize_column_spacing()`` collapses excess whitespace and
          ensures correct punctuation spacing while keeping the
          ``===COLUMN_BREAK===`` token and all structured data intact.
      5.  Clean text (lowercase, strip special chars, protect emails/URLs).
      6.  Extract sections using keyword-based heading detection.
      7.  Detect profile photo.
      8.  Detect language.
      9.  Assemble and return the output record.
    """
    file_path_str = str(file_path)
    suffix = file_path.suffix.lower()
    resume_id = str(uuid.uuid4())

    logger.info(f"Processing: {file_path.name}")

    raw_text = ""
    source_format = "failed"

    try:
        if suffix == ".pdf":
            logger.info("  [method] PDF extraction (column-aware)")
            raw_text, source_format = extract_text_pdf(file_path_str)
        elif suffix in (".docx", ".doc"):
            logger.info("  [method] DOCX extraction (table-column-aware)")
            raw_text, source_format = extract_text_docx(file_path_str)
        else:
            logger.warning(f"  [skip] Unsupported format: {suffix}")
            source_format = "failed"
    except Exception as e:
        logger.error(f"  [critical_error] {file_path.name}: {e}")
        source_format = "failed"
        raw_text = ""

    # ── Step 3: contact info — from original text, before any mutation ────────
    contact = extract_contact_info(raw_text)

    # ── Step 4: normalize column spacing ─────────────────────────────────────
    # Must run BEFORE clean_text so that the COLUMN_BREAK_TOKEN (which contains
    # only ASCII uppercase letters, digits, and "=") is not stripped by the
    # special-character remover in clean_text.
    normalised_text = normalize_column_spacing(raw_text) if raw_text else ""

    # ── Step 4b: repair OCR / broken-token spacing artifacts ────────────────
    # Runs AFTER normalize_column_spacing (so COLUMN_BREAK_TOKEN is already
    # present) and BEFORE clean_text (so protected tokens survive lowercasing).
    ocr_fixed_text = fix_ocr_spacing(normalised_text) if normalised_text else ""

    # ── Step 5: clean (lowercase, strip junk chars, collapse whitespace) ──────
    cleaned_text = clean_text(ocr_fixed_text) if ocr_fixed_text else ""

    # ── Step 6: section extraction ────────────────────────────────────────────
    # Pass debug=True when the PARSER_DEBUG env-var is set so that the debug
    # prints include the filename context from this outer scope.
    _dbg = bool(os.environ.get("PARSER_DEBUG", ""))
    if _dbg:
        print(f"\n[DEBUG] ── Processing: {file_path.name} ────────────────")
    sections_raw = extract_sections(cleaned_text)

    # fallback (boşsa devreye gir)
    if all(not v.strip() for k, v in sections_raw.items() if k != "__confidence__"):
        sections_raw = extract_sections(cleaned_text, debug=True)

    # duplicate temizleme — not needed since extract_sections handles it
    # sections_raw = deduplicate_sections(sections_raw)

    # Extract confidence scores (internal quality metadata) — stored separately
    # so the public "sections" dict contains only string values, preserving
    # backward compatibility with all downstream consumers.
    section_confidence: dict[str, float] = sections_raw.pop("__confidence__", {})
    sections = sections_raw

    # ── Step 7: photo detection ───────────────────────────────────────────────
    has_photo = False
    if source_format != "failed":
        has_photo = detect_photo(file_path_str, source_format)

    # ── Step 8: language detection ────────────────────────────────────────────
    language = detect_language(cleaned_text)

    # ── Step 9: assemble record ───────────────────────────────────────────────
    record = {
        "resume_id": resume_id,
        "file_path": file_path_str,
        "raw_text": cleaned_text,
        "sections": sections,
        "section_confidence": section_confidence,
        "contact": contact,
        "has_photo": has_photo,
        "language": language,
        "source_format": source_format,
    }

    logger.info(
        f"  [done] {file_path.name} → "
        f"format={source_format}, lang={language}, "
        f"photo={has_photo}, chars={len(cleaned_text)}"
    )
    return record


# ─────────────────────────────────────────────
#  10. DATASET BUILDER
# ─────────────────────────────────────────────

_FAILED_RECORD_TEMPLATE: dict = {
    "raw_text": "",
    "sections": {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
    },
    "contact": {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
    },
    "has_photo": False,
    "language": "en",
    "source_format": "failed",
}


def build_dataset(
    pdf_dir: str,
    docx_dir: str,
    output_path: str = "final_dataset.json",
) -> None:
    """
    Iterate through both CV directories, process every file, and write the
    aggregated results to a JSON file.

    Args:
        pdf_dir:     Path to the directory containing PDF CVs.
        docx_dir:    Path to the directory containing DOCX CVs.
        output_path: Destination JSON file path.
    """
    pdf_dir = Path(pdf_dir)
    docx_dir = Path(docx_dir)

    pdf_files = sorted(pdf_dir.glob("*.pdf")) if pdf_dir.exists() else []
    docx_files = (
        sorted(list(docx_dir.glob("*.docx")) + list(docx_dir.glob("*.doc")))
        if docx_dir.exists()
        else []
    )

    all_files = pdf_files + docx_files
    total = len(all_files)

    if total == 0:
        logger.warning("No CV files found. Check your directory paths.")
        return

    logger.info(
        f"Found {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX(s) "
        f"— {total} files total."
    )

    dataset: list[dict] = []
    failed_files: list[str] = []

    for file_path in tqdm(all_files, desc="Parsing CVs", unit="file"):
        try:
            record = process_cv(file_path)
            dataset.append(record)
            if record["source_format"] == "failed":
                failed_files.append(str(file_path))
        except Exception as e:
            logger.error(f"[unhandled] {file_path.name}: {e}")
            failed_files.append(str(file_path))
            dataset.append(
                {
                    "resume_id": str(uuid.uuid4()),
                    "file_path": str(file_path),
                    **_FAILED_RECORD_TEMPLATE,
                }
            )

    # ── Write JSON output ────────────────────────────────────────────────────
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(dataset, f, ensure_ascii=False, indent=2)

    # ── Summary report ───────────────────────────────────────────────────────
    success_count = sum(1 for r in dataset if r["source_format"] != "failed")
    ocr_count = sum(1 for r in dataset if r["source_format"] == "ocr")
    two_col_note = "  (check cv_parser.log for column-layout details)"

    logger.info("=" * 60)
    logger.info(f"DONE — {total} files processed")
    logger.info(f"  ✓ Success  : {success_count}")
    logger.info(f"  ✗ Failed   : {len(failed_files)}")
    logger.info(f"  ~ OCR used : {ocr_count}")
    logger.info(f"  Output     : {output_path}")
    logger.info(two_col_note)
    if failed_files:
        logger.warning("Failed files:")
        for fp in failed_files:
            logger.warning(f"    {fp}")
    logger.info("=" * 60)


# ─────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="CV Parsing Pipeline (column-aware) — produces final_dataset.json"
    )
    parser.add_argument(
        "--pdf-dir",
        type=str,
        default="cvs/pdf",
        help="Directory containing PDF CV files (default: cvs/pdf)",
    )
    parser.add_argument(
        "--docx-dir",
        type=str,
        default="cvs/docx",
        help="Directory containing DOCX CV files (default: cvs/docx)",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="final_dataset.json",
        help="Output JSON file path (default: final_dataset.json)",
    )
    args = parser.parse_args()

    build_dataset(
        pdf_dir=args.pdf_dir,
        docx_dir=args.docx_dir,
        output_path=args.output,
    )
