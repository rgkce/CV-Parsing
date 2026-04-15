"""
cv_parser.py
============
Production-level CV/Resume Parsing Pipeline
Processes PDF and DOCX files from two directories and outputs a structured JSON dataset.

Dependencies:
    pip install pdfplumber pymupdf python-docx pytesseract pillow langdetect tqdm

System dependency:
    Tesseract OCR must be installed: https://github.com/tesseract-ocr/tesseract
    On Ubuntu/Debian: sudo apt-get install tesseract-ocr tesseract-ocr-tur
    On macOS: brew install tesseract
"""

import os
import re
import json
import uuid
import logging
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
import pytesseract
from PIL import Image
import io
from tqdm import tqdm

try:
    from langdetect import detect as langdetect_detect

    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False


# ─────────────────────────────────────────────
#  LOGGING SETUP
# ─────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(levelname)s]  %(message)s",
    handlers=[
        logging.FileHandler("cv_parser.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  CONSTANTS / KEYWORD MAPS
# ─────────────────────────────────────────────

# OCR fallback threshold: if extracted text has fewer characters than this,
# we consider extraction a failure and invoke OCR.
OCR_FALLBACK_THRESHOLD = 80

# Section heading keywords — English and Turkish
SECTION_KEYWORDS: dict[str, list[str]] = {
    "summary": [
        "summary",
        "profile",
        "about",
        "about me",
        "objective",
        "professional summary",
        "career objective",
        "özet",
        "hakkımda",
        "hakkında",
        "profil",
        "kişisel özet",
    ],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
        "work history",
        "career history",
        "positions held",
        "deneyim",
        "iş deneyimi",
        "çalışma geçmişi",
        "kariyer",
    ],
    "education": [
        "education",
        "academic background",
        "academic history",
        "qualifications",
        "degrees",
        "schooling",
        "eğitim",
        "öğrenim",
        "akademik geçmiş",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "competencies",
        "technologies",
        "tools",
        "proficiencies",
        "key skills",
        "areas of expertise",
        "yetenekler",
        "beceriler",
        "teknolojiler",
        "yetkinlikler",
        "teknik beceriler",
    ],
    "projects": [
        "projects",
        "personal projects",
        "key projects",
        "portfolio",
        "open source",
        "projeler",
        "kişisel projeler",
    ],
}

# Turkish word list used for quick language heuristic
TURKISH_WORDS = {
    "ve",
    "bir",
    "bu",
    "da",
    "de",
    "için",
    "ile",
    "ben",
    "ama",
    "olan",
    "gibi",
    "çok",
    "ise",
    "ya",
    "ya da",
    "ki",
    "mı",
    "deneyim",
    "eğitim",
    "beceri",
    "hakkımda",
    "özet",
    "çalışma",
    "proje",
    "yetenek",
}


# ─────────────────────────────────────────────
#  1. TEXT EXTRACTION — PDF
# ─────────────────────────────────────────────


def _is_two_column_layout(page) -> bool:
    """
    Heuristic to detect two-column layouts using pdfplumber word bounding boxes.

    Strategy: If more than 30% of text words cluster in two distinct horizontal
    bands (left-half vs right-half of the page), we flag it as two-column.
    """
    words = page.extract_words()
    if not words:
        return False

    page_width = page.width
    mid = page_width / 2

    left_count = sum(1 for w in words if w["x1"] < mid)
    right_count = sum(1 for w in words if w["x0"] >= mid)
    total = len(words)

    if total == 0:
        return False

    # Both columns must have at least 20% of total words to qualify
    left_ratio = left_count / total
    right_ratio = right_count / total
    return left_ratio > 0.20 and right_ratio > 0.20


def _extract_two_column_text(page) -> str:
    """
    Reconstruct reading order for a two-column PDF page.

    Approach:
      1. Collect all word bounding boxes.
      2. Split words into LEFT and RIGHT columns by page midpoint.
      3. Sort each column top-to-bottom.
      4. Concatenate: left column text, then right column text.
      5. Group words into lines by proximity on the Y axis.
    """
    words = page.extract_words()
    if not words:
        return ""

    page_width = page.width
    mid = page_width / 2

    left_words = [w for w in words if w["x1"] < mid]
    right_words = [w for w in words if w["x0"] >= mid]

    def words_to_text(word_list: list) -> str:
        if not word_list:
            return ""
        # Sort by top-to-bottom, then left-to-right
        word_list.sort(key=lambda w: (round(w["top"], 1), w["x0"]))
        lines = []
        current_line: list[str] = []
        current_top: Optional[float] = None

        for w in word_list:
            if current_top is None or abs(w["top"] - current_top) < 5:
                current_line.append(w["text"])
                current_top = w["top"]
            else:
                lines.append(" ".join(current_line))
                current_line = [w["text"]]
                current_top = w["top"]

        if current_line:
            lines.append(" ".join(current_line))

        return "\n".join(lines)

    left_text = words_to_text(left_words)
    right_text = words_to_text(right_words)

    # Merge: left column first, then right column
    parts = [p for p in [left_text, right_text] if p.strip()]
    return "\n\n".join(parts)


def extract_text_pdf(file_path: str) -> tuple[str, str]:
    """
    Extract text from a PDF file.

    Returns:
        (text, source_format) where source_format is "pdf" or "ocr".

    Strategy:
      1. Try pdfplumber for standard text extraction.
      2. Detect two-column layouts per page — if found, reconstruct reading order.
      3. If extracted text is too short, fall back to OCR via PyMuPDF rasterisation.
    """
    file_path = str(file_path)
    all_pages_text: list[str] = []
    layout_issues: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    if _is_two_column_layout(page):
                        logger.info(
                            f"  [layout] Two-column layout detected — "
                            f"page {page_num} of '{os.path.basename(file_path)}'"
                        )
                        page_text = _extract_two_column_text(page)
                    else:
                        # Single-column: standard extraction, preserve natural reading order
                        page_text = (
                            page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                        )

                    all_pages_text.append(page_text)

                except Exception as page_err:
                    logger.warning(
                        f"  [layout_issue] page {page_num} of "
                        f"'{os.path.basename(file_path)}': {page_err}"
                    )
                    layout_issues.append(f"page_{page_num}")
                    all_pages_text.append("")

        full_text = "\n\n".join(filter(None, all_pages_text))

        if len(full_text.strip()) >= OCR_FALLBACK_THRESHOLD:
            if layout_issues:
                logger.warning(
                    f"  [layout_issue] '{os.path.basename(file_path)}' — "
                    f"problematic pages: {layout_issues}"
                )
            return full_text, "pdf"

        # Text is too short — fall through to OCR
        logger.info(
            f"  [pdf→ocr] Text too short ({len(full_text.strip())} chars) "
            f"in '{os.path.basename(file_path)}' — invoking OCR."
        )

    except Exception as e:
        logger.warning(
            f"  [pdf_error] pdfplumber failed on "
            f"'{os.path.basename(file_path)}': {e} — invoking OCR."
        )

    return ocr_fallback(file_path)


# ─────────────────────────────────────────────
#  2. TEXT EXTRACTION — DOCX
# ─────────────────────────────────────────────


def extract_text_docx(file_path: str) -> tuple[str, str]:
    """
    Extract text from a DOCX file.

    Handles:
      - Regular paragraphs
      - Tables (cell-by-cell, row-by-row — avoids merging unrelated columns)
      - Text boxes (shapes that contain text frames in the XML)

    Returns:
        (text, "docx")
    """
    file_path = str(file_path)
    doc = Document(file_path)
    parts: list[str] = []

    # --- Paragraphs (main body) ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- Tables ---
    # Each table cell is treated as an independent text block.
    # We do NOT concatenate adjacent cells on the same row because in
    # two-column CV tables, each column belongs to a different section.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    # --- Text boxes (drawing canvas shapes) ---
    # These are embedded as <w:drawing> → <wp:inline/anchor> → <a:t> elements.
    # We walk the XML directly because python-docx doesn't expose them natively.
    try:
        body = doc.element.body
        ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        for textbox in body.iter(qn("w:txbxContent")):
            for child_para in textbox.iter(qn("w:p")):
                texts = [node.text for node in child_para.iter(qn("w:t")) if node.text]
                if texts:
                    parts.append(" ".join(texts))
    except Exception as e:
        logger.warning(
            f"  [docx_textbox] Could not extract text boxes from "
            f"'{os.path.basename(file_path)}': {e}"
        )

    full_text = "\n".join(parts)
    logger.info(
        f"  [docx] Extracted {len(full_text)} chars from "
        f"'{os.path.basename(file_path)}'"
    )
    return full_text, "docx"


# ─────────────────────────────────────────────
#  3. OCR FALLBACK
# ─────────────────────────────────────────────


def ocr_fallback(file_path: str) -> tuple[str, str]:
    """
    Rasterise each page of a PDF with PyMuPDF and run Tesseract OCR.

    We try English + Turkish language packs (eng+tur).
    Falls back to English-only if the combined pack is unavailable.

    Returns:
        (text, "ocr") or ("", "failed") on complete failure.
    """
    file_path = str(file_path)
    all_text: list[str] = []

    try:
        pdf_doc = fitz.open(file_path)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            # Render at 300 DPI for good OCR accuracy
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))

            # Attempt combined language OCR, fall back to English
            for lang in ("eng+tur", "eng"):
                try:
                    text = pytesseract.image_to_string(img, lang=lang, config="--psm 6")
                    all_text.append(text)
                    break
                except pytesseract.pytesseract.TesseractError:
                    if lang == "eng":
                        logger.warning(
                            f"  [ocr_warning] Tesseract failed on page "
                            f"{page_num + 1} of '{os.path.basename(file_path)}'"
                        )

        pdf_doc.close()
        full_text = "\n\n".join(filter(None, all_text))
        logger.info(
            f"  [ocr] Extracted {len(full_text)} chars from "
            f"'{os.path.basename(file_path)}'"
        )
        return full_text, "ocr"

    except Exception as e:
        logger.error(
            f"  [ocr_failed] OCR failed for '{os.path.basename(file_path)}': {e}"
        )
        return "", "failed"


# ─────────────────────────────────────────────
#  4. TEXT CLEANING
# ─────────────────────────────────────────────

# Pre-compiled patterns for efficiency
_RE_EMAIL = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    re.IGNORECASE,
)
_RE_URL = re.compile(
    r"https?://[^\s]+|www\.[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+",
    re.IGNORECASE,
)
_RE_PHONE = re.compile(
    r"(?:\+?\d[\d\s\-().]{6,}\d)",
)
_RE_MULTI_SPACE = re.compile(r"[ \t]{2,}")
_RE_MULTI_NEWLINE = re.compile(r"\n{3,}")
_RE_SPECIAL_CHARS = re.compile(r"[^\w\s@.,:;()\-+/#&'\"/\\]", re.UNICODE)


def clean_text(text: str) -> str:
    """
    Selective text cleaning that preserves structured data.

    Rules:
      1. Replace protected tokens (emails, URLs, phone numbers) with placeholders.
      2. Strip excessive special characters.
      3. Normalize whitespace and newlines.
      4. Lowercase everything.
      5. Restore protected tokens.

    This ensures we never break email addresses, URLs, or phone numbers.
    """
    if not text:
        return ""

    # ── Step 1: Replace protected tokens with placeholders ──────────────────
    protected: dict[str, str] = {}

    def protect(pattern: re.Pattern, prefix: str, t: str) -> str:
        def replacer(m):
            key = f"__PROTECTED_{prefix}_{len(protected)}__"
            protected[key] = m.group(0)
            return key

        return pattern.sub(replacer, t)

    text = protect(_RE_EMAIL, "EMAIL", text)
    text = protect(_RE_URL, "URL", text)
    text = protect(_RE_PHONE, "PHONE", text)

    # ── Step 2: Remove unwanted special characters ───────────────────────────
    # Keep: word chars, spaces, common punctuation, unicode letters
    text = _RE_SPECIAL_CHARS.sub(" ", text)

    # ── Step 3: Normalize whitespace ─────────────────────────────────────────
    text = _RE_MULTI_SPACE.sub(" ", text)
    text = _RE_MULTI_NEWLINE.sub("\n\n", text)
    text = text.strip()

    # ── Step 4: Lowercase ─────────────────────────────────────────────────────
    text = text.lower()

    # ── Step 5: Restore protected tokens ─────────────────────────────────────
    for key, original in protected.items():
        text = text.replace(key.lower(), original)  # keys got lowercased in step 4

    return text


# ─────────────────────────────────────────────
#  5. SECTION DETECTION
# ─────────────────────────────────────────────


def _build_section_pattern() -> re.Pattern:
    """
    Build a compiled regex that matches section headings.

    The pattern accounts for:
      - Headings on their own line (possibly with trailing colons or spaces)
      - Case-insensitive matching (Turkish characters handled by UNICODE flag)
      - Optional surrounding whitespace
    """
    all_keywords = sorted(
        (kw for keywords in SECTION_KEYWORDS.values() for kw in keywords),
        key=len,
        reverse=True,  # match longer phrases first
    )
    escaped = [re.escape(kw) for kw in all_keywords]
    pattern = r"(?:^|\n)\s*(?:" + "|".join(escaped) + r")\s*:?\s*(?=\n|$)"
    return re.compile(pattern, re.IGNORECASE | re.MULTILINE | re.UNICODE)


_SECTION_HEADING_RE = _build_section_pattern()


def _classify_heading(heading: str) -> Optional[str]:
    """
    Given a detected heading string, return the canonical section name or None.
    """
    heading_lower = heading.strip().lower().rstrip(":")
    for section, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in heading_lower or heading_lower in kw:
                return section
    return None


def extract_sections(text: str) -> dict[str, str]:
    """
    Split the raw text into structured sections using keyword-based heading detection.

    Algorithm:
      1. Find all heading matches and their positions.
      2. Classify each heading into a canonical section.
      3. Slice the text between consecutive headings.
      4. Assign each slice to the corresponding section.
      5. If a section appears multiple times, concatenate (avoid overwrite).

    Returns a dict with keys: summary, experience, education, skills, projects.
    All missing sections default to "".
    """
    sections: dict[str, str] = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
    }

    # Find all heading matches
    matches = list(_SECTION_HEADING_RE.finditer(text))
    if not matches:
        return sections

    # Build list of (position, heading_text, canonical_section)
    headings: list[tuple[int, str, str]] = []
    for m in matches:
        canonical = _classify_heading(m.group(0))
        if canonical:
            headings.append((m.end(), m.group(0), canonical))

    # Extract text block between consecutive headings
    for idx, (start_pos, heading_text, section_name) in enumerate(headings):
        end_pos = headings[idx + 1][0] if idx + 1 < len(headings) else len(text)
        block = text[start_pos:end_pos].strip()
        # Concatenate if section already has content (handles split sections)
        if sections[section_name]:
            sections[section_name] += "\n\n" + block
        else:
            sections[section_name] = block

    return sections


# ─────────────────────────────────────────────
#  6. CONTACT INFO EXTRACTION
# ─────────────────────────────────────────────

# LinkedIn: match profile URLs with or without https://
_RE_LINKEDIN = re.compile(
    r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_\-%.]+)",
    re.IGNORECASE,
)

# GitHub: match profile URLs with or without https://
_RE_GITHUB = re.compile(
    r"(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_\-]+)",
    re.IGNORECASE,
)

# Phone: international format (+90 5xx xxx xx xx) or local, with spaces/dashes
_RE_PHONE_CONTACT = re.compile(
    r"(?<!\d)(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?){1,4}\d{2,4}(?!\d)",
)


def extract_contact_info(text: str) -> dict[str, str]:
    """
    Extract contact information using precise regular expressions.

    Extracts: email, phone, linkedin, github.

    IMPORTANT: We operate on the ORIGINAL (un-cleaned) text to avoid
    mangling punctuation inside emails and URLs.
    """
    contact: dict[str, str] = {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
    }

    # Email
    email_match = _RE_EMAIL.search(text)
    if email_match:
        contact["email"] = email_match.group(0).strip()

    # Phone — filter out obviously wrong matches (too short or too long)
    phone_matches = _RE_PHONE_CONTACT.findall(text)
    for raw in phone_matches:
        digits = re.sub(r"\D", "", raw)
        if 7 <= len(digits) <= 15:
            contact["phone"] = raw.strip()
            break

    # LinkedIn — return the full URL (reconstruct if needed)
    linkedin_match = _RE_LINKEDIN.search(text)
    if linkedin_match:
        full = linkedin_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["linkedin"] = full

    # GitHub — same treatment
    github_match = _RE_GITHUB.search(text)
    if github_match:
        full = github_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["github"] = full

    return contact


# ─────────────────────────────────────────────
#  7. PHOTO DETECTION
# ─────────────────────────────────────────────


def detect_photo_pdf(file_path: str) -> bool:
    """
    Detect whether the PDF contains at least one embedded image (likely a profile photo).

    Uses PyMuPDF's get_images() — returns all images per page.
    We ignore tiny images (< 50×50 px) which are likely icons or decorations.
    """
    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            images = page.get_images(full=True)
            for img in images:
                # img tuple: (xref, smask, width, height, bpc, colorspace, ...)
                width = img[2]
                height = img[3]
                if width >= 50 and height >= 50:
                    doc.close()
                    return True
        doc.close()
    except Exception as e:
        logger.warning(
            f"  [photo_pdf] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo_docx(file_path: str) -> bool:
    """
    Detect whether a DOCX file contains inline images (likely a profile photo).

    Checks:
      - Inline shapes with pictures
      - Relationships of type image in the document part
    """
    try:
        doc = Document(str(file_path))

        # Check inline shapes
        for shape in doc.inline_shapes:
            # InlineShape type 3 = picture
            if shape.type is not None:
                return True

        # Also check document relationships for any image parts
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                return True

    except Exception as e:
        logger.warning(
            f"  [photo_docx] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo(file_path: str, source_format: str) -> bool:
    """
    Dispatch to format-specific photo detection.
    """
    if source_format in ("pdf", "ocr"):
        return detect_photo_pdf(file_path)
    elif source_format == "docx":
        return detect_photo_docx(file_path)
    return False


# ─────────────────────────────────────────────
#  8. LANGUAGE DETECTION
# ─────────────────────────────────────────────


def detect_language(text: str) -> str:
    """
    Detect whether the CV is in Turkish, English, or mixed.

    Strategy:
      1. Simple heuristic: count Turkish stopwords in lowercased text.
         If > 3 Turkish words found → likely Turkish or mixed.
      2. If langdetect is available: use it for a two-pass decision.
         - Sample first 1000 chars (usually the header/summary — most diagnostic)
         - Sample last 500 chars
         - If both agree → that language
         - If they disagree → "mixed"
      3. Fallback: if only Turkish words heuristic → use that.

    Returns "tr" | "en" | "mixed"
    """
    if not text or len(text.strip()) < 30:
        return "en"

    sample = text[:2000].lower()
    words_in_sample = set(re.findall(r"\b\w+\b", sample))
    turkish_hits = words_in_sample & TURKISH_WORDS

    if LANGDETECT_AVAILABLE:
        try:
            # Two samples for robustness
            sample_start = text[:1000]
            sample_end = text[-500:]

            lang_start = langdetect_detect(sample_start)
            lang_end = langdetect_detect(sample_end)

            # Normalize to tr/en
            def norm(lang: str) -> str:
                if lang == "tr":
                    return "tr"
                return "en"

            l1 = norm(lang_start)
            l2 = norm(lang_end)

            if l1 == l2:
                return l1
            else:
                return "mixed"

        except Exception:
            pass  # fall through to heuristic

    # Heuristic fallback
    if len(turkish_hits) >= 3:
        return "tr"
    return "en"


# ─────────────────────────────────────────────
#  9. MAIN PROCESSING FUNCTION
# ─────────────────────────────────────────────


def process_cv(file_path: Path) -> dict:
    """
    Process a single CV file (PDF or DOCX) and return a structured record.

    Steps:
      1. Determine file format.
      2. Extract raw text (with OCR fallback).
      3. Extract contact info from ORIGINAL text (before cleaning).
      4. Clean text.
      5. Detect sections.
      6. Detect photo.
      7. Detect language.
      8. Build and return the output record.
    """
    file_path_str = str(file_path)
    suffix = file_path.suffix.lower()
    resume_id = str(uuid.uuid4())

    logger.info(f"Processing: {file_path.name}")

    # ── 1. Determine format and extract text ─────────────────────────────────
    raw_text = ""
    source_format = "failed"

    try:
        if suffix == ".pdf":
            logger.info(f"  [method] PDF extraction")
            raw_text, source_format = extract_text_pdf(file_path_str)
        elif suffix in (".docx", ".doc"):
            logger.info(f"  [method] DOCX extraction")
            raw_text, source_format = extract_text_docx(file_path_str)
        else:
            logger.warning(f"  [skip] Unsupported format: {suffix}")
            source_format = "failed"
    except Exception as e:
        logger.error(f"  [critical_error] {file_path.name}: {e}")
        source_format = "failed"
        raw_text = ""

    # ── 2. Extract contact info from original (uncleaned) text ───────────────
    contact = extract_contact_info(raw_text)

    # ── 3. Clean text ─────────────────────────────────────────────────────────
    cleaned_text = clean_text(raw_text) if raw_text else ""

    # ── 4. Extract sections ───────────────────────────────────────────────────
    sections = extract_sections(cleaned_text)

    # ── 5. Detect photo ───────────────────────────────────────────────────────
    has_photo = False
    if source_format != "failed":
        has_photo = detect_photo(file_path_str, source_format)

    # ── 6. Detect language ────────────────────────────────────────────────────
    language = detect_language(cleaned_text)

    # ── 7. Assemble record ────────────────────────────────────────────────────
    record = {
        "resume_id": resume_id,
        "file_path": file_path_str,
        "raw_text": cleaned_text,
        "sections": sections,
        "contact": contact,
        "has_photo": has_photo,
        "language": language,
        "source_format": source_format,
    }

    logger.info(
        f"  [done] {file_path.name} → "
        f"format={source_format}, lang={language}, "
        f"photo={has_photo}, chars={len(cleaned_text)}"
    )
    return record


# ─────────────────────────────────────────────
#  10. DATASET BUILDER
# ─────────────────────────────────────────────


def build_dataset(
    pdf_dir: str,
    docx_dir: str,
    output_path: str = "final_dataset.json",
) -> None:
    """
    Iterate through both CV directories, process every file, and write the
    aggregated results to a JSON file.

    Args:
        pdf_dir:     Path to the directory containing PDF CVs.
        docx_dir:    Path to the directory containing DOCX CVs.
        output_path: Destination JSON file path.
    """
    pdf_dir = Path(pdf_dir)
    docx_dir = Path(docx_dir)

    # Collect all files — sorted for deterministic ordering
    pdf_files = sorted(pdf_dir.glob("*.pdf")) if pdf_dir.exists() else []
    docx_files = (
        sorted(list(docx_dir.glob("*.docx")) + list(docx_dir.glob("*.doc")))
        if docx_dir.exists()
        else []
    )

    all_files = pdf_files + docx_files
    total = len(all_files)

    if total == 0:
        logger.warning("No CV files found. Check your directory paths.")
        return

    logger.info(
        f"Found {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX(s) "
        f"— {total} files total."
    )

    dataset: list[dict] = []
    failed_files: list[str] = []

    for file_path in tqdm(all_files, desc="Parsing CVs", unit="file"):
        try:
            record = process_cv(file_path)
            dataset.append(record)
            if record["source_format"] == "failed":
                failed_files.append(str(file_path))
        except Exception as e:
            logger.error(f"[unhandled] {file_path.name}: {e}")
            failed_files.append(str(file_path))
            # Still add a minimal failed record to keep IDs consistent
            dataset.append(
                {
                    "resume_id": str(uuid.uuid4()),
                    "file_path": str(file_path),
                    "raw_text": "",
                    "sections": {
                        "summary": "",
                        "experience": "",
                        "education": "",
                        "skills": "",
                        "projects": "",
                    },
                    "contact": {
                        "email": "",
                        "phone": "",
                        "linkedin": "",
                        "github": "",
                    },
                    "has_photo": False,
                    "language": "en",
                    "source_format": "failed",
                }
            )

    # ── Write JSON output ─────────────────────────────────────────────────────
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(dataset, f, ensure_ascii=False, indent=2)

    # ── Summary report ────────────────────────────────────────────────────────
    success_count = sum(1 for r in dataset if r["source_format"] != "failed")
    ocr_count = sum(1 for r in dataset if r["source_format"] == "ocr")

    logger.info("=" * 60)
    logger.info(f"DONE — {total} files processed")
    logger.info(f"  ✓ Success  : {success_count}")
    logger.info(f"  ✗ Failed   : {len(failed_files)}")
    logger.info(f"  ~ OCR used : {ocr_count}")
    logger.info(f"  Output     : {output_path}")
    if failed_files:
        logger.warning("Failed files:")
        for fp in failed_files:
            logger.warning(f"    {fp}")
    logger.info("=" * 60)


# ─────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="CV Parsing Pipeline — produces final_dataset.json"
    )
    parser.add_argument(
        "--pdf-dir",
        type=str,
        default="cvs/pdf",
        help="Directory containing PDF CV files (default: cvs/pdf)",
    )
    parser.add_argument(
        "--docx-dir",
        type=str,
        default="cvs/docx",
        help="Directory containing DOCX CV files (default: cvs/docx)",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="final_dataset.json",
        help="Output JSON file path (default: final_dataset.json)",
    )
    args = parser.parse_args()

    build_dataset(
        pdf_dir=args.pdf_dir,
        docx_dir=args.docx_dir,
        output_path=args.output,
    )
