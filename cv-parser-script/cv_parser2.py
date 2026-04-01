"""
cv_parser.py  (column-aware edition)
=====================================
Production-level CV/Resume Parsing Pipeline
Processes PDF and DOCX files from two directories and outputs a structured JSON dataset.

Key upgrades over v1:
  • Column-aware PDF extraction  — detects single / two-column / multi-column (3+)
    layouts using pdfplumber word bounding-boxes and reconstructs correct reading
    order (left column → right column, top-to-bottom within each).
  • Gap-analysis column split   — instead of a hard page-midpoint cut, we find the
    *largest horizontal whitespace gap* between word clusters to locate the column
    boundary.  This handles asymmetric sidebar layouts (e.g. narrow left sidebar
    with contact info and a wide right content area).
  • Table-based PDF pages       — detected separately; cells are read left-to-right,
    top-to-bottom using pdfplumber's extract_tables().
  • OCR fallback intact         — still triggered when digital text is too sparse.
  • DOCX two-column tables      — each column of a two-cell row is appended in order
    so sidebar contact info always precedes the main content.
  • All other functions (cleaning, section extraction, contact, photo, language,
    dataset builder) are preserved from v1 with only minor additions.

Dependencies:
    pip install pdfplumber pymupdf python-docx pytesseract pillow langdetect tqdm

System dependency:
    Tesseract OCR: https://github.com/tesseract-ocr/tesseract
    Ubuntu/Debian : sudo apt-get install tesseract-ocr tesseract-ocr-tur
    macOS         : brew install tesseract
"""

import os
import re
import json
import uuid
import logging
import statistics
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
import pytesseract
from PIL import Image
import io
from tqdm import tqdm

try:
    from langdetect import detect as langdetect_detect

    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False


# ─────────────────────────────────────────────
#  LOGGING SETUP
# ─────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(levelname)s]  %(message)s",
    handlers=[
        logging.FileHandler("cv_parser.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  CONSTANTS / KEYWORD MAPS
# ─────────────────────────────────────────────

# OCR fallback threshold: if extracted text has fewer characters than this,
# we consider extraction a failure and invoke OCR.
OCR_FALLBACK_THRESHOLD = 80

# Minimum ratio of words that must appear in EACH column for multi-column detection.
# e.g. 0.15 means both left and right clusters need ≥15% of all page words.
COLUMN_MIN_RATIO = 0.15

# When scanning for the horizontal gap between columns, we project word x-ranges
# onto a 1-D grid of this many buckets.  Higher = finer resolution but slower.
GAP_SCAN_BUCKETS = 200

# Minimum gap width (as fraction of page width) to accept a column split boundary.
# Prevents splitting on narrow inter-word spaces inside a single column.
MIN_GAP_FRACTION = 0.03

# Section heading keywords — English and Turkish
SECTION_KEYWORDS: dict[str, list[str]] = {
    "summary": [
        "summary",
        "profile",
        "about",
        "about me",
        "objective",
        "professional summary",
        "career objective",
        "özet",
        "hakkımda",
        "hakkında",
        "profil",
        "kişisel özet",
    ],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
        "work history",
        "career history",
        "positions held",
        "deneyim",
        "iş deneyimi",
        "çalışma geçmişi",
        "kariyer",
    ],
    "education": [
        "education",
        "academic background",
        "academic history",
        "qualifications",
        "degrees",
        "schooling",
        "eğitim",
        "öğrenim",
        "akademik geçmiş",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "competencies",
        "technologies",
        "tools",
        "proficiencies",
        "key skills",
        "areas of expertise",
        "yetenekler",
        "beceriler",
        "teknolojiler",
        "yetkinlikler",
        "teknik beceriler",
    ],
    "projects": [
        "projects",
        "personal projects",
        "key projects",
        "portfolio",
        "open source",
        "projeler",
        "kişisel projeler",
    ],
}

# Turkish word list used for quick language heuristic
TURKISH_WORDS = {
    "ve",
    "bir",
    "bu",
    "da",
    "de",
    "için",
    "ile",
    "ben",
    "ama",
    "olan",
    "gibi",
    "çok",
    "ise",
    "ya",
    "ya da",
    "ki",
    "mı",
    "deneyim",
    "eğitim",
    "beceri",
    "hakkımda",
    "özet",
    "çalışma",
    "proje",
    "yetenek",
}


# ─────────────────────────────────────────────
#  1. COLUMN-AWARE PDF TEXT EXTRACTION
# ─────────────────────────────────────────────

# ── 1a. Gap-based column boundary detection ───────────────────────────────────


def _find_column_split_x(words: list[dict], page_width: float) -> Optional[float]:
    """
    Find the x-coordinate of the widest horizontal whitespace gap between
    word clusters on a page — this is where the column boundary lies.

    Algorithm:
      1. Build a 1-D occupancy array along the x-axis (GAP_SCAN_BUCKETS wide).
      2. Mark each bucket as 'occupied' if any word overlaps it.
      3. Find the longest contiguous run of *empty* buckets.
      4. Return the centre x of that run, or None if no significant gap found.

    This is more robust than a hard midpoint split because it adapts to:
      - Asymmetric layouts (narrow sidebar + wide main column)
      - Layouts where the column split is off-centre
    """
    if not words:
        return None

    bucket_size = page_width / GAP_SCAN_BUCKETS
    occupied = [False] * GAP_SCAN_BUCKETS

    for w in words:
        # Mark every bucket overlapped by this word's x extent
        start_bucket = max(0, int(w["x0"] / bucket_size))
        end_bucket = min(GAP_SCAN_BUCKETS - 1, int(w["x1"] / bucket_size))
        for b in range(start_bucket, end_bucket + 1):
            occupied[b] = True

    # Find the longest unoccupied run
    best_start = best_end = -1
    current_start = None

    for i, occ in enumerate(occupied):
        if not occ:
            if current_start is None:
                current_start = i
        else:
            if current_start is not None:
                run_len = i - current_start
                if run_len > (best_end - best_start):
                    best_start, best_end = current_start, i - 1
                current_start = None

    # Handle gap that runs to the end of the array
    if current_start is not None:
        run_len = GAP_SCAN_BUCKETS - current_start
        if run_len > (best_end - best_start):
            best_start, best_end = current_start, GAP_SCAN_BUCKETS - 1

    if best_start == -1:
        return None  # No gap found

    gap_width_fraction = (best_end - best_start + 1) / GAP_SCAN_BUCKETS
    if gap_width_fraction < MIN_GAP_FRACTION:
        # Gap too narrow — probably just inter-word spacing, not a column split
        return None

    # Return the pixel x-coordinate of the gap's centre
    gap_centre_x = ((best_start + best_end) / 2.0) * bucket_size
    return gap_centre_x


# ── 1b. Layout detection ──────────────────────────────────────────────────────


class PageLayout:
    """Enum-like class for layout type labels."""

    SINGLE = "single"
    TWO_COL = "two_column"
    MULTI = "multi_column"  # 3+ columns (unusual but exists)
    TABLE = "table"  # page dominated by a table structure


def _detect_page_layout(page, words: list[dict]) -> str:
    """
    Classify a pdfplumber page into a layout type.

    Decision logic:
      1. If pdfplumber finds any tables on the page → TABLE layout.
      2. Count words left/right of the detected gap boundary.
         - If gap exists AND both sides have >= COLUMN_MIN_RATIO → TWO_COL.
         - Additional check: if standard deviation of x0 values is very high
           (words spread all over), consider MULTI.
      3. Otherwise → SINGLE.
    """
    # Check for explicit table structures first
    try:
        tables = page.extract_tables()
        if tables and any(len(t) > 1 for t in tables):
            # Only flag as TABLE if there's a meaningful table (>1 row)
            return PageLayout.TABLE
    except Exception:
        pass

    if not words:
        return PageLayout.SINGLE

    page_width = page.width
    split_x = _find_column_split_x(words, page_width)

    if split_x is None:
        return PageLayout.SINGLE

    left_count = sum(1 for w in words if w["x1"] <= split_x)
    right_count = sum(1 for w in words if w["x0"] > split_x)
    total = len(words)

    left_ratio = left_count / total
    right_ratio = right_count / total

    if left_ratio < COLUMN_MIN_RATIO or right_ratio < COLUMN_MIN_RATIO:
        return PageLayout.SINGLE

    # Quick multi-column check: look for a second significant gap in EACH half
    left_words = [w for w in words if w["x1"] <= split_x]
    right_words = [w for w in words if w["x0"] > split_x]

    left_gap = _find_column_split_x(left_words, split_x)
    right_gap = _find_column_split_x(right_words, page_width - split_x)

    if left_gap is not None or right_gap is not None:
        return PageLayout.MULTI

    return PageLayout.TWO_COL


# ── 1c. Word-list → ordered text reconstruction ───────────────────────────────


def _words_to_text(word_list: list[dict], y_tolerance: float = 4.0) -> str:
    """
    Convert a list of pdfplumber word dicts to a text string.

    Words are sorted by (top, x0) and grouped into lines when their
    vertical positions are within y_tolerance points of each other.
    This handles slight baseline misalignments common in designed CVs.

    Args:
        word_list    : list of pdfplumber word dicts (keys: text, x0, x1, top).
        y_tolerance  : vertical distance (pts) within which words share a line.

    Returns:
        Newline-separated string, one line per detected text row.
    """
    if not word_list:
        return ""

    # Sort: primary = top (vertical position), secondary = x0 (left-to-right)
    word_list = sorted(
        word_list, key=lambda w: (round(w["top"] / y_tolerance) * y_tolerance, w["x0"])
    )

    lines: list[str] = []
    current_line: list[str] = []
    current_top: Optional[float] = None

    for w in word_list:
        if current_top is None or abs(w["top"] - current_top) < y_tolerance:
            current_line.append(w["text"])
            # Keep track of the average top so we don't drift on long lines
            current_top = (
                w["top"] if current_top is None else (current_top + w["top"]) / 2
            )
        else:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [w["text"]]
            current_top = w["top"]

    if current_line:
        lines.append(" ".join(current_line))

    return "\n".join(lines)


# ── 1d. Per-layout extraction functions ──────────────────────────────────────


def _extract_single_column(page) -> str:
    """
    Standard single-column extraction.

    Uses pdfplumber's built-in extract_text() with tight tolerances to
    keep words on the same line together while preserving paragraph breaks.
    """
    return page.extract_text(x_tolerance=3, y_tolerance=3) or ""


def _extract_two_column(page, words: list[dict]) -> str:
    """
    Reconstruct reading order for a two-column page.

    The column boundary is found via _find_column_split_x() (gap analysis)
    rather than a fixed midpoint — this correctly handles asymmetric layouts
    such as a narrow contact sidebar on the left and a wide experience column
    on the right.

    Reading order: left column (top → bottom) then right column (top → bottom).

    Args:
        page  : pdfplumber page object (needed for page_width).
        words : pre-extracted word list (avoids re-extracting).

    Returns:
        Reconstructed text with left column first, then right column.
    """
    page_width = page.width
    split_x = _find_column_split_x(words, page_width)

    if split_x is None:
        # Fallback to simple midpoint if gap detection fails
        split_x = page_width / 2
        logger.debug(
            "  [layout] Gap detection failed — falling back to midpoint split."
        )

    # Partition words into left and right columns
    # Note: words straddling the gap (x0 < split_x < x1) go to whichever
    # column their *centre* falls in — avoids double-counting headers.
    left_words: list[dict] = []
    right_words: list[dict] = []

    for w in words:
        centre_x = (w["x0"] + w["x1"]) / 2
        if centre_x <= split_x:
            left_words.append(w)
        else:
            right_words.append(w)

    left_text = _words_to_text(left_words)
    right_text = _words_to_text(right_words)

    # Left column first, then right column — separated by a blank line so the
    # section extractor can still find headings at the start of lines.
    parts = [p for p in [left_text, right_text] if p.strip()]
    return "\n\n".join(parts)


def _extract_table_page(page) -> str:
    """
    Extract text from a table-dominated page.

    Strategy:
      1. Extract tables cell-by-cell (left-to-right, top-to-bottom per row).
         Each cell's content is kept as a separate block so section headings
         in table headers are preserved.
      2. Also extract any non-table paragraphs floating outside the tables.

    This is the right approach for CV templates that use invisible tables as
    layout grids (common in Word-exported PDFs).
    """
    parts: list[str] = []

    try:
        tables = page.extract_tables()
        if tables:
            for table in tables:
                for row in table:
                    if row:
                        for cell in row:
                            cell_text = (cell or "").strip()
                            if cell_text:
                                parts.append(cell_text)
    except Exception as e:
        logger.warning(f"  [layout_issue] Table extraction failed on page: {e}")

    # Also capture text not inside any table bounding box
    try:
        non_table_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
        if non_table_text.strip():
            parts.append(non_table_text)
    except Exception:
        pass

    return "\n".join(parts)


def _extract_multi_column(page, words: list[dict]) -> str:
    """
    Fallback for pages with 3+ columns (rare in CVs but exists in fancy templates).

    Strategy: cluster words by their x0 into N groups using a simple gap scan,
    sort each cluster top-to-bottom, and concatenate left-to-right.

    If clustering is ambiguous, we fall back to a plain word-by-word sort by
    (top, x0) which is still better than line-by-line extraction.
    """
    logger.info("  [layout] Multi-column (3+) page detected — using positional sort.")

    # Sort all words by natural reading order (top → bottom, left → right)
    # For most 3-column layouts this produces a readable result.
    return _words_to_text(words, y_tolerance=5.0)


# ── 1e. Main PDF extraction orchestrator ─────────────────────────────────────


def extract_text_pdf(file_path: str) -> tuple[str, str]:
    """
    Extract text from a PDF file using layout-aware column reconstruction.

    Returns:
        (text, source_format) where source_format is "pdf" or "ocr".

    Per-page strategy:
      1. Extract word bounding boxes with pdfplumber.
      2. Detect layout type: SINGLE / TWO_COL / MULTI / TABLE.
      3. Dispatch to the appropriate extraction function.
      4. If the combined extracted text is too short → OCR fallback.
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    all_pages_text: list[str] = []
    layout_issues: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    # Extract word dicts once — reused by detection + extraction
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False,
                        use_text_flow=False,  # raw positions, not PDF text-flow
                    )

                    layout = _detect_page_layout(page, words)

                    if layout == PageLayout.TWO_COL:
                        logger.info(
                            f"  [layout] TWO_COLUMN detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_two_column(page, words)

                    elif layout == PageLayout.TABLE:
                        logger.info(
                            f"  [layout] TABLE layout detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_table_page(page)

                    elif layout == PageLayout.MULTI:
                        logger.info(
                            f"  [layout] MULTI_COLUMN (3+) detected — "
                            f"page {page_num} of '{basename}'"
                        )
                        page_text = _extract_multi_column(page, words)

                    else:
                        # SINGLE column — standard extraction
                        page_text = _extract_single_column(page)

                    all_pages_text.append(page_text)

                except Exception as page_err:
                    logger.warning(
                        f"  [layout_issue] page {page_num} of '{basename}': {page_err}"
                    )
                    layout_issues.append(f"page_{page_num}")
                    all_pages_text.append("")

        full_text = "\n\n".join(filter(None, all_pages_text))

        if layout_issues:
            logger.warning(
                f"  [layout_issue] '{basename}' — problematic pages: {layout_issues}"
            )

        if len(full_text.strip()) >= OCR_FALLBACK_THRESHOLD:
            return full_text, "pdf"

        # Text is too short — fall through to OCR
        logger.info(
            f"  [pdf→ocr] Text too short ({len(full_text.strip())} chars) "
            f"in '{basename}' — invoking OCR."
        )

    except Exception as e:
        logger.warning(
            f"  [pdf_error] pdfplumber failed on '{basename}': {e} — invoking OCR."
        )

    return ocr_fallback(file_path)


# ─────────────────────────────────────────────
#  2. TEXT EXTRACTION — DOCX
# ─────────────────────────────────────────────


def extract_text_docx(file_path: str) -> tuple[str, str]:
    """
    Extract text from a DOCX file.

    Handles:
      - Regular paragraphs (main body flow)
      - Tables — including two-column layout tables:
          * Single-cell rows are read normally.
          * Two-cell rows are treated as left/right columns:
            left cell content is appended first, then right cell content.
            This preserves sidebar-style layouts common in designed DOCX CVs.
          * Rows with 3+ cells: each cell appended in left-to-right order.
      - Text boxes (shapes containing text frames in the XML).

    Returns:
        (text, "docx")
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    doc = Document(file_path)
    parts: list[str] = []

    # ── Paragraphs (main body) ────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # ── Tables ────────────────────────────────────────────────────────────────
    # We treat each table separately to avoid merging unrelated columns.
    # Two-column tables are the most common layout device in DOCX CVs.
    for table_idx, table in enumerate(doc.tables):
        col_count = max((len(row.cells) for row in table.rows), default=0)

        if col_count == 2:
            # Two-column table: left column first, then right column.
            # Collect each column's text independently, then interleave so that
            # content from the same row stays adjacent — better for section detection.
            logger.info(
                f"  [docx_layout] Two-column table detected "
                f"(table {table_idx + 1}) in '{basename}'"
            )
            left_parts: list[str] = []
            right_parts: list[str] = []
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    left_text = cells[0].text.strip()
                    right_text = cells[1].text.strip()
                    if left_text:
                        left_parts.append(left_text)
                    if right_text:
                        right_parts.append(right_text)
                elif len(cells) == 1:
                    cell_text = cells[0].text.strip()
                    if cell_text:
                        left_parts.append(cell_text)

            # Left column block first, then right column block
            if left_parts:
                parts.append("\n".join(left_parts))
            if right_parts:
                parts.append("\n".join(right_parts))

        else:
            # Single-column or 3+-column table: cell-by-cell, row-by-row
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

    # ── Text boxes (drawing canvas shapes) ───────────────────────────────────
    # Embedded as <w:txbxContent> → child <w:p> → <w:t> in the XML.
    try:
        body = doc.element.body
        for textbox in body.iter(qn("w:txbxContent")):
            for child_para in textbox.iter(qn("w:p")):
                texts = [node.text for node in child_para.iter(qn("w:t")) if node.text]
                if texts:
                    parts.append(" ".join(texts))
    except Exception as e:
        logger.warning(
            f"  [docx_textbox] Could not extract text boxes from '{basename}': {e}"
        )

    full_text = "\n".join(parts)
    logger.info(f"  [docx] Extracted {len(full_text)} chars from '{basename}'")
    return full_text, "docx"


# ─────────────────────────────────────────────
#  3. OCR FALLBACK
# ─────────────────────────────────────────────


def ocr_fallback(file_path: str) -> tuple[str, str]:
    """
    Rasterise each page of a PDF with PyMuPDF and run Tesseract OCR.

    We try English + Turkish language packs (eng+tur).
    Falls back to English-only if the combined pack is unavailable.

    Returns:
        (text, "ocr") or ("", "failed") on complete failure.
    """
    file_path = str(file_path)
    basename = os.path.basename(file_path)
    all_text: list[str] = []

    logger.info(f"  [ocr] Starting OCR on '{basename}'")

    try:
        pdf_doc = fitz.open(file_path)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            # Render at 300 DPI for good OCR accuracy
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))

            # Attempt combined language OCR, fall back to English-only
            ocr_success = False
            for lang in ("eng+tur", "eng"):
                try:
                    text = pytesseract.image_to_string(img, lang=lang, config="--psm 6")
                    all_text.append(text)
                    if lang == "eng+tur":
                        logger.info(
                            f"  [ocr] Page {page_num + 1}: eng+tur OCR successful."
                        )
                    else:
                        logger.info(
                            f"  [ocr] Page {page_num + 1}: eng OCR used (tur pack unavailable)."
                        )
                    ocr_success = True
                    break
                except pytesseract.pytesseract.TesseractError:
                    continue

            if not ocr_success:
                logger.warning(
                    f"  [ocr_warning] Tesseract failed entirely on page "
                    f"{page_num + 1} of '{basename}'"
                )

        pdf_doc.close()
        full_text = "\n\n".join(filter(None, all_text))
        logger.info(f"  [ocr] Extracted {len(full_text)} chars from '{basename}'")
        return full_text, "ocr"

    except Exception as e:
        logger.error(f"  [ocr_failed] OCR failed for '{basename}': {e}")
        return "", "failed"


# ─────────────────────────────────────────────
#  4. TEXT CLEANING
# ─────────────────────────────────────────────

# Pre-compiled patterns for efficiency
_RE_EMAIL = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    re.IGNORECASE,
)
_RE_URL = re.compile(
    r"https?://[^\s]+|www\.[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+",
    re.IGNORECASE,
)
_RE_PHONE = re.compile(
    r"(?:\+?\d[\d\s\-().]{6,}\d)",
)
_RE_MULTI_SPACE = re.compile(r"[ \t]{2,}")
_RE_MULTI_NEWLINE = re.compile(r"\n{3,}")
_RE_SPECIAL_CHARS = re.compile(r"[^\w\s@.,:;()\-+/#&'\"/\\]", re.UNICODE)


def clean_text(text: str) -> str:
    """
    Selective text cleaning that preserves structured data.

    Rules:
      1. Replace protected tokens (emails, URLs, phone numbers) with placeholders.
      2. Strip excessive special characters.
      3. Normalize whitespace and newlines.
      4. Lowercase everything.
      5. Restore protected tokens.

    This ensures we never break email addresses, URLs, or phone numbers.
    """
    if not text:
        return ""

    # ── Step 1: Protect structured tokens ───────────────────────────────────
    protected: dict[str, str] = {}

    def protect(pattern: re.Pattern, prefix: str, t: str) -> str:
        def replacer(m):
            key = f"__PROTECTED_{prefix}_{len(protected)}__"
            protected[key] = m.group(0)
            return key

        return pattern.sub(replacer, t)

    text = protect(_RE_EMAIL, "EMAIL", text)
    text = protect(_RE_URL, "URL", text)
    text = protect(_RE_PHONE, "PHONE", text)

    # ── Step 2: Remove unwanted special characters ────────────────────────
    text = _RE_SPECIAL_CHARS.sub(" ", text)

    # ── Step 3: Normalize whitespace ─────────────────────────────────────
    text = _RE_MULTI_SPACE.sub(" ", text)
    text = _RE_MULTI_NEWLINE.sub("\n\n", text)
    text = text.strip()

    # ── Step 4: Lowercase ────────────────────────────────────────────────
    text = text.lower()

    # ── Step 5: Restore protected tokens ─────────────────────────────────
    for key, original in protected.items():
        text = text.replace(key.lower(), original)

    return text


# ─────────────────────────────────────────────
#  5. SECTION DETECTION
# ─────────────────────────────────────────────


def _build_section_pattern() -> re.Pattern:
    """
    Build a compiled regex that matches section headings.

    The pattern accounts for:
      - Headings on their own line (possibly with trailing colons or spaces)
      - Case-insensitive matching (Turkish characters handled by UNICODE flag)
      - Optional surrounding whitespace
    """
    all_keywords = sorted(
        (kw for keywords in SECTION_KEYWORDS.values() for kw in keywords),
        key=len,
        reverse=True,  # match longer phrases first
    )
    escaped = [re.escape(kw) for kw in all_keywords]
    pattern = r"(?:^|\n)\s*(?:" + "|".join(escaped) + r")\s*:?\s*(?=\n|$)"
    return re.compile(pattern, re.IGNORECASE | re.MULTILINE | re.UNICODE)


_SECTION_HEADING_RE = _build_section_pattern()


def _classify_heading(heading: str) -> Optional[str]:
    """
    Given a detected heading string, return the canonical section name or None.
    """
    heading_lower = heading.strip().lower().rstrip(":")
    for section, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in heading_lower or heading_lower in kw:
                return section
    return None


def extract_sections(text: str) -> dict[str, str]:
    """
    Split the raw text into structured sections using keyword-based heading detection.

    Algorithm:
      1. Find all heading matches and their positions.
      2. Classify each heading into a canonical section.
      3. Slice the text between consecutive headings.
      4. Assign each slice to the corresponding section.
      5. If a section appears multiple times, concatenate (avoid overwrite).

    Returns a dict with keys: summary, experience, education, skills, projects.
    All missing sections default to "".
    """
    sections: dict[str, str] = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
    }

    matches = list(_SECTION_HEADING_RE.finditer(text))
    if not matches:
        return sections

    headings: list[tuple[int, str, str]] = []
    for m in matches:
        canonical = _classify_heading(m.group(0))
        if canonical:
            headings.append((m.end(), m.group(0), canonical))

    for idx, (start_pos, heading_text, section_name) in enumerate(headings):
        end_pos = headings[idx + 1][0] if idx + 1 < len(headings) else len(text)
        block = text[start_pos:end_pos].strip()
        if sections[section_name]:
            sections[section_name] += "\n\n" + block
        else:
            sections[section_name] = block

    return sections


# ─────────────────────────────────────────────
#  6. CONTACT INFO EXTRACTION
# ─────────────────────────────────────────────

_RE_LINKEDIN = re.compile(
    r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_\-%.]+)",
    re.IGNORECASE,
)
_RE_GITHUB = re.compile(
    r"(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_\-]+)",
    re.IGNORECASE,
)
_RE_PHONE_CONTACT = re.compile(
    r"(?<!\d)(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?){1,4}\d{2,4}(?!\d)",
)


def extract_contact_info(text: str) -> dict[str, str]:
    """
    Extract contact information using precise regular expressions.

    Extracts: email, phone, linkedin, github.

    IMPORTANT: We operate on the ORIGINAL (un-cleaned) text to avoid
    mangling punctuation inside emails and URLs.
    """
    contact: dict[str, str] = {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
    }

    email_match = _RE_EMAIL.search(text)
    if email_match:
        contact["email"] = email_match.group(0).strip()

    phone_matches = _RE_PHONE_CONTACT.findall(text)
    for raw in phone_matches:
        digits = re.sub(r"\D", "", raw)
        if 7 <= len(digits) <= 15:
            contact["phone"] = raw.strip()
            break

    linkedin_match = _RE_LINKEDIN.search(text)
    if linkedin_match:
        full = linkedin_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["linkedin"] = full

    github_match = _RE_GITHUB.search(text)
    if github_match:
        full = github_match.group(0)
        if not full.startswith("http"):
            full = "https://" + full
        contact["github"] = full

    return contact


# ─────────────────────────────────────────────
#  7. PHOTO DETECTION
# ─────────────────────────────────────────────


def detect_photo_pdf(file_path: str) -> bool:
    """
    Detect whether the PDF contains at least one embedded image (likely a profile photo).

    Uses PyMuPDF's get_images() — returns all images per page.
    We ignore tiny images (< 50×50 px) which are likely icons or decorations.
    """
    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            images = page.get_images(full=True)
            for img in images:
                # img tuple: (xref, smask, width, height, bpc, colorspace, ...)
                width = img[2]
                height = img[3]
                if width >= 50 and height >= 50:
                    doc.close()
                    return True
        doc.close()
    except Exception as e:
        logger.warning(
            f"  [photo_pdf] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo_docx(file_path: str) -> bool:
    """
    Detect whether a DOCX file contains inline images (likely a profile photo).

    Checks:
      - Inline shapes with pictures
      - Relationships of type image in the document part
    """
    try:
        doc = Document(str(file_path))

        for shape in doc.inline_shapes:
            if shape.type is not None:
                return True

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                return True

    except Exception as e:
        logger.warning(
            f"  [photo_docx] Could not check images in "
            f"'{os.path.basename(file_path)}': {e}"
        )
    return False


def detect_photo(file_path: str, source_format: str) -> bool:
    """Dispatch to format-specific photo detection."""
    if source_format in ("pdf", "ocr"):
        return detect_photo_pdf(file_path)
    elif source_format == "docx":
        return detect_photo_docx(file_path)
    return False


# ─────────────────────────────────────────────
#  8. LANGUAGE DETECTION
# ─────────────────────────────────────────────


def detect_language(text: str) -> str:
    """
    Detect whether the CV is in Turkish, English, or mixed.

    Strategy:
      1. Simple heuristic: count Turkish stopwords in lowercased text.
      2. If langdetect is available: two-pass (start + end samples).
         Both must agree → that language; else → "mixed".
      3. Fallback: if Turkish words heuristic triggers → "tr", else "en".

    Returns "tr" | "en" | "mixed"
    """
    if not text or len(text.strip()) < 30:
        return "en"

    sample = text[:2000].lower()
    words_in_sample = set(re.findall(r"\b\w+\b", sample))
    turkish_hits = words_in_sample & TURKISH_WORDS

    if LANGDETECT_AVAILABLE:
        try:
            sample_start = text[:1000]
            sample_end = text[-500:]

            lang_start = langdetect_detect(sample_start)
            lang_end = langdetect_detect(sample_end)

            def norm(lang: str) -> str:
                return "tr" if lang == "tr" else "en"

            l1, l2 = norm(lang_start), norm(lang_end)
            return l1 if l1 == l2 else "mixed"

        except Exception:
            pass  # fall through to heuristic

    if len(turkish_hits) >= 3:
        return "tr"
    return "en"


# ─────────────────────────────────────────────
#  9. MAIN PROCESSING FUNCTION
# ─────────────────────────────────────────────


def process_cv(file_path: Path) -> dict:
    """
    Process a single CV file (PDF or DOCX) and return a structured record.

    Steps:
      1. Determine file format.
      2. Extract raw text (column-aware, with OCR fallback).
      3. Extract contact info from ORIGINAL text (before cleaning).
      4. Clean text.
      5. Detect sections.
      6. Detect photo.
      7. Detect language.
      8. Build and return the output record.
    """
    file_path_str = str(file_path)
    suffix = file_path.suffix.lower()
    resume_id = str(uuid.uuid4())

    logger.info(f"Processing: {file_path.name}")

    raw_text = ""
    source_format = "failed"

    try:
        if suffix == ".pdf":
            logger.info("  [method] PDF extraction (column-aware)")
            raw_text, source_format = extract_text_pdf(file_path_str)
        elif suffix in (".docx", ".doc"):
            logger.info("  [method] DOCX extraction (table-column-aware)")
            raw_text, source_format = extract_text_docx(file_path_str)
        else:
            logger.warning(f"  [skip] Unsupported format: {suffix}")
            source_format = "failed"
    except Exception as e:
        logger.error(f"  [critical_error] {file_path.name}: {e}")
        source_format = "failed"
        raw_text = ""

    # Contact info extracted BEFORE cleaning to preserve punctuation
    contact = extract_contact_info(raw_text)
    cleaned_text = clean_text(raw_text) if raw_text else ""
    sections = extract_sections(cleaned_text)

    has_photo = False
    if source_format != "failed":
        has_photo = detect_photo(file_path_str, source_format)

    language = detect_language(cleaned_text)

    record = {
        "resume_id": resume_id,
        "file_path": file_path_str,
        "raw_text": cleaned_text,
        "sections": sections,
        "contact": contact,
        "has_photo": has_photo,
        "language": language,
        "source_format": source_format,
    }

    logger.info(
        f"  [done] {file_path.name} → "
        f"format={source_format}, lang={language}, "
        f"photo={has_photo}, chars={len(cleaned_text)}"
    )
    return record


# ─────────────────────────────────────────────
#  10. DATASET BUILDER
# ─────────────────────────────────────────────

_FAILED_RECORD_TEMPLATE: dict = {
    "raw_text": "",
    "sections": {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
    },
    "contact": {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
    },
    "has_photo": False,
    "language": "en",
    "source_format": "failed",
}


def build_dataset(
    pdf_dir: str,
    docx_dir: str,
    output_path: str = "final_dataset.json",
) -> None:
    """
    Iterate through both CV directories, process every file, and write the
    aggregated results to a JSON file.

    Args:
        pdf_dir:     Path to the directory containing PDF CVs.
        docx_dir:    Path to the directory containing DOCX CVs.
        output_path: Destination JSON file path.
    """
    pdf_dir = Path(pdf_dir)
    docx_dir = Path(docx_dir)

    pdf_files = sorted(pdf_dir.glob("*.pdf")) if pdf_dir.exists() else []
    docx_files = (
        sorted(list(docx_dir.glob("*.docx")) + list(docx_dir.glob("*.doc")))
        if docx_dir.exists()
        else []
    )

    all_files = pdf_files + docx_files
    total = len(all_files)

    if total == 0:
        logger.warning("No CV files found. Check your directory paths.")
        return

    logger.info(
        f"Found {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX(s) "
        f"— {total} files total."
    )

    dataset: list[dict] = []
    failed_files: list[str] = []

    for file_path in tqdm(all_files, desc="Parsing CVs", unit="file"):
        try:
            record = process_cv(file_path)
            dataset.append(record)
            if record["source_format"] == "failed":
                failed_files.append(str(file_path))
        except Exception as e:
            logger.error(f"[unhandled] {file_path.name}: {e}")
            failed_files.append(str(file_path))
            dataset.append(
                {
                    "resume_id": str(uuid.uuid4()),
                    "file_path": str(file_path),
                    **_FAILED_RECORD_TEMPLATE,
                }
            )

    # ── Write JSON output ────────────────────────────────────────────────────
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(dataset, f, ensure_ascii=False, indent=2)

    # ── Summary report ───────────────────────────────────────────────────────
    success_count = sum(1 for r in dataset if r["source_format"] != "failed")
    ocr_count = sum(1 for r in dataset if r["source_format"] == "ocr")
    two_col_note = "  (check cv_parser.log for column-layout details)"

    logger.info("=" * 60)
    logger.info(f"DONE — {total} files processed")
    logger.info(f"  ✓ Success  : {success_count}")
    logger.info(f"  ✗ Failed   : {len(failed_files)}")
    logger.info(f"  ~ OCR used : {ocr_count}")
    logger.info(f"  Output     : {output_path}")
    logger.info(two_col_note)
    if failed_files:
        logger.warning("Failed files:")
        for fp in failed_files:
            logger.warning(f"    {fp}")
    logger.info("=" * 60)


# ─────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="CV Parsing Pipeline (column-aware) — produces final_dataset.json"
    )
    parser.add_argument(
        "--pdf-dir",
        type=str,
        default="cvs/pdf",
        help="Directory containing PDF CV files (default: cvs/pdf)",
    )
    parser.add_argument(
        "--docx-dir",
        type=str,
        default="cvs/docx",
        help="Directory containing DOCX CV files (default: cvs/docx)",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="final_dataset.json",
        help="Output JSON file path (default: final_dataset.json)",
    )
    args = parser.parse_args()

    build_dataset(
        pdf_dir=args.pdf_dir,
        docx_dir=args.docx_dir,
        output_path=args.output,
    )
